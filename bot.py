import os
import re
import logging
import asyncio
import random
import time
from io import BytesIO
from math import ceil
from typing import List, Dict, Optional, Tuple
from contextlib import asynccontextmanager
import hashlib

from fastapi import FastAPI, Request, HTTPException

from docx import Document
from pypdf import PdfReader

import database_quiz as db
import ai_parser
import pro_features
import parser_engine
from telegram import (
    Update,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    BotCommand,
    BotCommandScopeDefault,
    BotCommandScopeAllPrivateChats,
    BotCommandScopeAllGroupChats,
    BotCommandScopeChatMember,
)
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    PollAnswerHandler,
    ContextTypes,
    filters,
)

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)

TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")

# Telegram webhook configuration.
# No long-polling/getUpdates process is used in this version.
RENDER_EXTERNAL_URL = (os.environ.get("RENDER_EXTERNAL_URL") or "").strip().rstrip("/")
RENDER_EXTERNAL_HOSTNAME = (os.environ.get("RENDER_EXTERNAL_HOSTNAME") or "").strip()
WEBHOOK_BASE_URL = (os.environ.get("WEBHOOK_BASE_URL") or "").strip().rstrip("/")

if not WEBHOOK_BASE_URL:
    if RENDER_EXTERNAL_URL:
        WEBHOOK_BASE_URL = RENDER_EXTERNAL_URL
    elif RENDER_EXTERNAL_HOSTNAME:
        WEBHOOK_BASE_URL = f"https://{RENDER_EXTERNAL_HOSTNAME}"
    else:
        WEBHOOK_BASE_URL = "https://quizbot-py-1.onrender.com"

WEBHOOK_PATH = "/telegram"
WEBHOOK_URL = f"{WEBHOOK_BASE_URL}{WEBHOOK_PATH}"

# Telegram signs webhook requests with this secret header.
# It rotates automatically if TELEGRAM_TOKEN changes.
WEBHOOK_SECRET = hashlib.sha256(
    (TELEGRAM_TOKEN or "missing-token").encode("utf-8")
).hexdigest()[:48]

TELEGRAM_APP: Optional[Application] = None

# Fast prototype storage.
# NOTE: This is kept in RAM, so a Render restart clears current sessions.
USER_DATA: Dict[int, dict] = {}
GROUP_DATA: Dict[int, dict] = {}
POLL_MAP: Dict[str, dict] = {}
WELCOME_SHOWN = set()

GROUP_SIZES = [30, 40, 50, 100]

# Telegram supports timed polls. These are the study choices shown before Start.
TIMER_CHOICES = [10, 15, 20, 30, 40, 60, 120]
GROUP_EMPTY_STOP_THRESHOLD = 3
GROUP_SETUP_TTL_SECONDS = 15 * 60
FREE_IMPORT_LIMIT = 1
PRO_AI_IMPORT_LIMIT = int(os.environ.get("PRO_AI_IMPORT_LIMIT", "20"))
MAX_UPLOAD_BYTES = int(os.environ.get("MAX_UPLOAD_BYTES", str(20 * 1024 * 1024)))
QUESTION_TRANSITION_DELAY = 1.0

OWNER_TELEGRAM_ID = int(os.environ.get("OWNER_TELEGRAM_ID", "0") or 0)

PRIVATE_COMMAND_MENU = [
    BotCommand("start", "Bosh menyuni ochish"),
    BotCommand("new", "Yangi test yuklash"),
    BotCommand("quizzes", "Saqlangan testlarim"),
    BotCommand("continue", "Oxirgi quizni davom ettirish"),
    BotCommand("pause", "Faol quizni vaqtincha pauza qilish"),
    BotCommand("resume", "Pauzadagi quizni davom ettirish"),
    BotCommand("stop", "Faol quizni to‘liq to‘xtatish"),
    BotCommand("progress", "Natijalarim"),
    BotCommand("group", "Guruh quiz rejimi"),
    BotCommand("settings", "Sozlamalar"),
    BotCommand("plan", "Tarif holati"),
    BotCommand("help", "Yordam"),
]

GROUP_COMMAND_MENU = [
    BotCommand("start", "Guruh menyusini ochish"),
    BotCommand("group", "Guruh quizini boshlash/boshqarish"),
    BotCommand("quizzes", "Mening saqlangan testlarim"),
    BotCommand("help", "Guruh quiz bo‘yicha yordam"),
]

GROUP_HOST_COMMAND_MENU = [
    BotCommand("group", "Guruh quiz boshqaruv paneli"),
    BotCommand("quizzes", "Mening saqlangan testlarim"),
    BotCommand("new", "Yangi test yuklash"),
    BotCommand("pause", "Guruh quizini vaqtincha pauza qilish"),
    BotCommand("resume", "Pauzadagi quizni davom ettirish"),
    BotCommand("skip", "Joriy savolni o‘tkazib yuborish"),
    BotCommand("stop", "Quizni to‘liq to‘xtatish"),
    BotCommand("release", "Quiz boshqaruvini bo‘shatish"),
    BotCommand("parser", "Joriy test parser hisobotini ko‘rish"),
    BotCommand("help", "Yordam"),
]


async def set_host_command_menu(bot, chat_id: int, user_id: int) -> None:
    """Show controller-only commands after '/' only to the current group host."""
    try:
        await bot.set_my_commands(
            GROUP_HOST_COMMAND_MENU,
            scope=BotCommandScopeChatMember(chat_id=chat_id, user_id=user_id),
        )
    except Exception:
        logging.exception("Could not set host command menu")


async def clear_host_command_menu(bot, chat_id: int, user_id: int) -> None:
    """Remove the host-specific menu so Telegram falls back to normal group commands."""
    try:
        await bot.delete_my_commands(
            scope=BotCommandScopeChatMember(chat_id=chat_id, user_id=user_id),
        )
    except Exception:
        logging.exception("Could not clear host command menu")


async def claim_group_host_with_menu(chat_id: int, user, bot) -> Tuple[bool, Optional[str]]:
    """
    Claim the group host lock and immediately give that user the controller command menu.
    If an old setup lock went stale, also clear its stale command menu.
    """
    previous = GROUP_DATA.get(chat_id)
    stale_previous_id = None
    if previous and host_is_stale(previous):
        stale_previous_id = previous.get("controller_id")

    ok, host_name = claim_group_host(chat_id, user)
    if not ok:
        return ok, host_name

    if stale_previous_id and stale_previous_id != user.id:
        await clear_host_command_menu(bot, chat_id, int(stale_previous_id))

    await set_host_command_menu(bot, chat_id, user.id)
    return True, None


def format_duration(seconds: int) -> str:
    if seconds == 60:
        return "1 daqiqa"
    if seconds == 120:
        return "2 daqiqa"
    return f"{seconds} soniya"


def is_owner(user_id: int) -> bool:
    return bool(OWNER_TELEGRAM_ID and user_id == OWNER_TELEGRAM_ID)


async def user_has_pro(user_id: int) -> bool:
    if is_owner(user_id):
        return True
    if not db.is_enabled():
        return False
    try:
        status = await db.get_plan_status(user_id, free_limit=FREE_IMPORT_LIMIT)
        return bool(status.get("is_pro"))
    except Exception:
        logging.exception("PRO holatini tekshirib bo‘lmadi")
        return False


async def effective_preferences(user_id: int) -> dict:
    if db.is_enabled():
        try:
            return await db.get_user_preferences(user_id)
        except Exception:
            logging.exception("Sozlamalarni o‘qib bo‘lmadi")
    return {
        "shuffle_questions": False,
        "shuffle_options": False,
        "quiz_mode": "practice",
        "default_group_size": 50,
        "default_timer": 30,
    }


def quiz_mode_label(mode: str) -> str:
    return "📝 Imtihon" if mode == "exam" else "📖 Mashq"


def prepare_poll_options(item: dict, shuffle_options: bool) -> Tuple[List[str], int]:
    """Return displayed options + the correct index after optional shuffling."""
    order = list(range(len(item["options"])))
    if shuffle_options:
        random.shuffle(order)
    displayed = [item["options"][i] for i in order]
    correct_original = int(item["correct_index"])
    correct_displayed = order.index(correct_original)
    return displayed, correct_displayed


def apply_question_order(session: dict) -> List[dict]:
    source = session.get("custom_questions") or session.get("questions", [])
    ordered = list(source)
    if session.get("shuffle_questions"):
        random.shuffle(ordered)
    return ordered


def host_is_stale(session: Optional[dict]) -> bool:
    if not session or session.get("active"):
        return False
    last = float(session.get("host_last_activity") or 0)
    return bool(last and (time.time() - last) > GROUP_SETUP_TTL_SECONDS)


def clean_stale_group_host(chat_id: int) -> None:
    session = GROUP_DATA.get(chat_id)
    if host_is_stale(session):
        GROUP_DATA.pop(chat_id, None)


def claim_group_host(chat_id: int, user) -> Tuple[bool, Optional[str]]:
    """One controller per Telegram group. Stale setup locks are released automatically."""
    clean_stale_group_host(chat_id)
    session = GROUP_DATA.get(chat_id)
    if session and session.get("controller_id") not in (None, user.id):
        return False, session.get("controller_name") or "boshqa foydalanuvchi"

    if not session:
        GROUP_DATA[chat_id] = {
            "chat_id": chat_id,
            "filename": None,
            "questions": [],
            "warnings": [],
            "group_size": None,
            "groups": [],
            "active": None,
            "results": {},
            "controller_id": user.id,
            "controller_name": user.full_name,
            "host_last_activity": time.time(),
            "last_leaderboard_text": None,
            "shuffle_questions": False,
            "shuffle_options": False,
        }
    else:
        session["controller_id"] = user.id
        session["controller_name"] = user.full_name
        session["host_last_activity"] = time.time()
    return True, None


def touch_group_host(session: dict) -> None:
    session["host_last_activity"] = time.time()


def release_group_host(chat_id: int) -> None:
    session = GROUP_DATA.get(chat_id)
    if not session:
        return
    if session.get("active"):
        return
    GROUP_DATA.pop(chat_id, None)


# -----------------------------
# TEXT EXTRACTION
# -----------------------------
def extract_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def extract_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Also read tables because many test files are stored in Word tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


# -----------------------------
# TEST PARSER
# -----------------------------
# The parser is deliberately tolerant about QUESTION NUMBER formatting.
# These all count as the same kind of question start:
#   № 123. Question
#   №123. Question
#   № 123.Question
#   № 123 Question
#   123. Question
#   123) Question
# Correct answers are still taken only from explicit answer information.
OPTION_RE = re.compile(
    r"^\s*([+*✓✔✅☑]?)\s*([A-Ha-h])\s*[\.\)\-:]\s*(.+?)\s*$"
)
ANSWER_LINE_RE = re.compile(
    r"^\s*(?:answer|correct\s*answer|javob|to['’`ʻ]?g['’`ʻ]?ri\s*javob)\s*[:\-]?\s*(.*?)\s*$",
    re.IGNORECASE,
)
ANSWER_KEY_PAIR_RE = re.compile(
    r"(?<!\d)(\d{1,4})\s*[\.\)\-:]?\s*([A-Ha-h])\b"
)


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    text = text.replace("–", "-").replace("—", "-")

    # IMPORTANT:
    # Do NOT globally split before patterns such as "D. ...". In teacher-made
    # documents this is often an author's initial (e.g. "Quronov D.") rather
    # than answer option D. Splitting it used to break otherwise valid questions.
    return text


def parse_question_start(line: str) -> Optional[Tuple[int, str]]:
    """Recognize a numbered question without caring about missing spaces/dots."""
    s = line.strip()

    # №-style numbering: punctuation after the number is optional.
    # Examples: №123, № 123., № 123.Question, № 123 Question
    m = re.match(r"^№\s*(\d{1,4})\s*[\.\)\-:]?\s*(.*)$", s, re.IGNORECASE)
    if m:
        return int(m.group(1)), m.group(2).strip()

    # Plain numbering: punctuation is expected, but spacing is optional.
    # Examples: 123.Question, 123. Question, 123) Question
    m = re.match(r"^(\d{1,4})\s*([\.\)\-:])\s*(.*)$", s)
    if m:
        return int(m.group(1)), m.group(3).strip()

    # A number on its own can also introduce a question on the next line.
    m = re.match(r"^(\d{1,4})\s*$", s)
    if m:
        return int(m.group(1)), ""

    return None


def is_source_header(text: str) -> bool:
    """Return True when the text after the question number is source metadata."""
    cleaned = re.sub(r"^\s*#\s*", "", text).strip().lower()
    return cleaned.startswith("manba") or cleaned.startswith("source")


def clean_question_line(line: str) -> str:
    """Remove harmless teacher markers without changing the actual question."""
    s = line.strip()
    if s in {"#", "=", "*", "+", "-"}:
        return ""
    return re.sub(r"^\s*#\s*", "", s).strip()


def parse_answer_key(text: str) -> Dict[int, str]:
    """
    Finds answer-key pairs such as:
      1-A
      2. B
      3) C
    Works best when the answer key is near the end of the file.
    """
    answers: Dict[int, str] = {}

    lower = text.lower()
    anchors = [
        "answer key",
        "answers",
        "javoblar",
        "javob kaliti",
        "to'g'ri javoblar",
        "to‘g‘ri javoblar",
    ]

    start = -1
    for anchor in anchors:
        pos = lower.rfind(anchor)
        if pos > start:
            start = pos

    if start < 0:
        return answers

    candidate = text[start:]
    for num, letter in ANSWER_KEY_PAIR_RE.findall(candidate):
        answers[int(num)] = letter.upper()

    return answers


def parse_questions(text: str) -> Tuple[List[dict], List[str]]:
    """
    Flexible question detection + strict answer validation.

    The parser tolerates missing spaces and punctuation around question numbers,
    but it NEVER guesses a correct answer. If the answer is missing or multiple
    answers are explicitly supplied, that block is reported as a warning.
    """
    text = normalize_text(text)
    answer_key = parse_answer_key(text)
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    questions: List[dict] = []
    warnings: List[str] = []

    current: Optional[dict] = None
    current_option_letter: Optional[str] = None

    def save_current():
        nonlocal current, current_option_letter
        if not current:
            return

        number = current["number"]
        options = current["options"]
        letter_to_index = current["letter_to_index"]

        answer_letters = list(current.get("answer_letters") or [])
        if not answer_letters and number in answer_key:
            answer_letters = [answer_key[number]]

        correct_index = None
        if len(answer_letters) == 1:
            correct_index = letter_to_index.get(answer_letters[0])

        reasons = []
        if not current["question"].strip():
            reasons.append("question text missing")
        if not (2 <= len(options) <= 10):
            reasons.append(f"{len(options)} options")
        if len(answer_letters) == 0:
            reasons.append("correct answer not found")
        elif len(answer_letters) > 1:
            reasons.append("multiple correct answers: " + ", ".join(answer_letters))
        elif correct_index is None:
            reasons.append(f"answer {answer_letters[0]} has no matching option")

        if not reasons:
            questions.append(
                {
                    "number": number,
                    "question": current["question"].strip(),
                    "options": [x.strip() for x in options],
                    "correct_index": correct_index,
                }
            )
        else:
            warnings.append(f"Question {number}: " + ", ".join(reasons))

        current = None
        current_option_letter = None

    for line in lines:
        q_start = parse_question_start(line)
        opt_match = OPTION_RE.match(line)
        ans_match = ANSWER_LINE_RE.match(line)

        # A numbered question start always wins over normal continuation text.
        # OPTION_RE is checked only to avoid treating something like "A) ..." as
        # a numeric question in unusual documents.
        if q_start and not opt_match:
            number, tail = q_start

            # Skip answer-key-style lines such as "1. B".
            if len(tail) == 1 and tail.upper() in "ABCDEFGH":
                continue

            save_current()
            current = {
                "number": number,
                # "№ 1. Manba: ..." is metadata; the real question follows.
                "question": "" if is_source_header(tail) else clean_question_line(tail),
                "options": [],
                "letter_to_index": {},
                "answer_letters": [],
            }
            current_option_letter = None
            continue

        if current is None:
            continue

        # Explicit answer line. Capture the WHOLE value so "Javob: B, D" is
        # recognized as ambiguous instead of silently choosing B.
        if ans_match:
            letters = [
                x.upper()
                for x in re.findall(r"\b([A-Ha-h])\b", ans_match.group(1))
            ]
            current["answer_letters"] = list(dict.fromkeys(letters))
            current_option_letter = None
            continue

        if opt_match:
            marker = opt_match.group(1)
            letter = opt_match.group(2).upper()
            option_text = opt_match.group(3).strip()

            # Conservative support for common teacher correct-answer markers:
            # +A) ..., *B) ..., ✓C) ..., ✔D) ...
            if marker:
                if letter not in current["answer_letters"]:
                    current["answer_letters"].append(letter)

            # Keep compatibility with the earlier "A) option *" convention.
            if option_text.endswith("*"):
                option_text = option_text[:-1].rstrip()
                if letter not in current["answer_letters"]:
                    current["answer_letters"].append(letter)

            current["letter_to_index"][letter] = len(current["options"])
            current["options"].append(option_text)
            current_option_letter = letter
            continue

        # Continuation lines. A leading # is a common teacher marker for the
        # question itself and should not become part of the visible Telegram text.
        if current_option_letter and current["options"]:
            current["options"][-1] += " " + line
        else:
            q_line = clean_question_line(line)
            if q_line:
                if current["question"]:
                    current["question"] += " " + q_line
                else:
                    current["question"] = q_line

    save_current()
    return questions, warnings


# -----------------------------
# TELEGRAM UI
# -----------------------------
def main_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton("📚 Testlarim", callback_data="menu_quizzes"),
                InlineKeyboardButton("📄 Yangi test", callback_data="menu_new"),
            ],
            [
                InlineKeyboardButton("▶️ Davom etish", callback_data="menu_continue"),
                InlineKeyboardButton("📊 Natijalar", callback_data="menu_progress"),
            ],
            [
                InlineKeyboardButton("👥 Guruh testi", callback_data="menu_group"),
                InlineKeyboardButton("⚙️ Sozlamalar", callback_data="menu_settings"),
            ],
            [
                InlineKeyboardButton("❓ Yordam", callback_data="menu_help"),
            ],
        ]
    )


def home_button() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        [[InlineKeyboardButton("🏠 Bosh menyu", callback_data="menu_home")]]
    )


def is_group_chat(chat) -> bool:
    return bool(chat and chat.type in ("group", "supergroup"))


def group_size_keyboard(
    group_mode: bool = False,
    default_size: Optional[int] = None,
) -> InlineKeyboardMarkup:
    prefix = "gsize" if group_mode else "size"
    custom_cb = "gcustomrange" if group_mode else "pcustomrange"
    random_cb = "gcustomrandom" if group_mode else "pcustomrandom"
    back_callback = "g_home" if group_mode else "menu_home"

    def label(size: int) -> str:
        return f"⭐ {size}" if default_size == size else str(size)

    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton(label(30), callback_data=f"{prefix}:30"),
                InlineKeyboardButton(label(40), callback_data=f"{prefix}:40"),
            ],
            [
                InlineKeyboardButton(label(50), callback_data=f"{prefix}:50"),
                InlineKeyboardButton(label(100), callback_data=f"{prefix}:100"),
            ],
            [InlineKeyboardButton("🎯 PRO: Oraliqni o‘zim tanlayman", callback_data=custom_cb)],
            [InlineKeyboardButton("🎲 PRO: Tasodifiy N ta savol", callback_data=random_cb)],
            [InlineKeyboardButton("🏠 Bosh menyu", callback_data=back_callback)],
        ]
    )


def timer_keyboard(
    group_index: int,
    group_mode: bool = False,
    default_timer: Optional[int] = None,
) -> InlineKeyboardMarkup:
    prefix = "gtimer" if group_mode else "timer"
    back_cb = "ggroups" if group_mode else "groups"
    back_text = "📚 Bo‘limlar" if group_mode else "📚 Guruhlar"

    def label(seconds: int) -> str:
        text = format_duration(seconds)
        return f"⭐ {text}" if default_timer == seconds else text

    return InlineKeyboardMarkup([
        [InlineKeyboardButton(label(10), callback_data=f"{prefix}:{group_index}:10"),
         InlineKeyboardButton(label(15), callback_data=f"{prefix}:{group_index}:15")],
        [InlineKeyboardButton(label(20), callback_data=f"{prefix}:{group_index}:20"),
         InlineKeyboardButton(label(30), callback_data=f"{prefix}:{group_index}:30")],
        [InlineKeyboardButton(label(40), callback_data=f"{prefix}:{group_index}:40"),
         InlineKeyboardButton(label(60), callback_data=f"{prefix}:{group_index}:60")],
        [InlineKeyboardButton(label(120), callback_data=f"{prefix}:{group_index}:120")],
        [InlineKeyboardButton(back_text, callback_data=back_cb)],
    ])


def group_home_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        [
            [
                InlineKeyboardButton("📚 Mening testlarim", callback_data="g_saved"),
                InlineKeyboardButton("📄 Test yuklash", callback_data="g_new"),
            ],
            [
                InlineKeyboardButton("📚 Joriy test", callback_data="g_current"),
                InlineKeyboardButton("🏆 Oxirgi reyting", callback_data="g_leaderboard"),
            ],
            [
                InlineKeyboardButton("🧪 Parser hisoboti", callback_data="g_parser"),
                InlineKeyboardButton("🔓 Sessiyani tugatish", callback_data="g_release"),
            ],
            [InlineKeyboardButton("❓ Yordam", callback_data="g_help")],
        ]
    )


def order_settings_keyboard(session: dict, group_mode: bool) -> InlineKeyboardMarkup:
    q_on = bool(session.get("shuffle_questions"))
    a_on = bool(session.get("shuffle_options"))
    if group_mode:
        q_cb, a_cb, done_cb, back_cb = "gtoq", "gtoa", "gorderdone", "g_home"
    else:
        q_cb, a_cb, done_cb, back_cb = "ptoq", "ptoa", "porderdone", "menu_home"
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(f"{'✅' if q_on else '⬜'} Savollarni aralashtirish", callback_data=q_cb)],
        [InlineKeyboardButton(f"{'✅' if a_on else '⬜'} Variantlarni aralashtirish", callback_data=a_cb)],
        [InlineKeyboardButton("➡️ Davom etish", callback_data=done_cb)],
        [InlineKeyboardButton("← Orqaga", callback_data=back_cb)],
    ])


async def send_order_settings(message, session: dict, group_mode: bool):
    await message.reply_text(
        "🔀 Quiz tartibini tanlang\n\n"
        "Savollarni aralashtirish — savollar boshqa tartibda beriladi.\n"
        "Variantlarni aralashtirish — A/B/C/D joylashuvi har savolda o‘zgaradi, "
        "lekin to‘g‘ri javob avtomatik moslashtiriladi.\n\n"
        "Test rejimi START dan oldin alohida tanlanadi.",
        reply_markup=order_settings_keyboard(session, group_mode),
    )


async def send_home(message):
    await message.reply_text(
        "🎓 Testchi\n\n"
        "Kerakli bo‘limni tanlang:",
        reply_markup=main_menu(),
    )


async def send_group_home(message):
    chat_id = message.chat.id
    clean_stale_group_host(chat_id)
    session = GROUP_DATA.get(chat_id)
    host_text = "👤 Boshqaruvchi: yo‘q"
    if session and session.get("controller_id"):
        host_text = f"👤 Boshqaruvchi: {session.get('controller_name') or 'foydalanuvchi'}"

    await message.reply_text(
        "👥 Testchi — guruh rejimi\n\n"
        f"{host_text}\n\n"
        "Quizni boshlash uchun /group yuboring yoki ‘Mening testlarim’/‘Test yuklash’ni bosing. "
        "Bir vaqtda faqat bitta foydalanuvchi quizni boshqaradi.\n\n"
        "Savol taymeri tugamaguncha keyingi savol chiqmaydi. "
        f"{GROUP_EMPTY_STOP_THRESHOLD} ta ketma-ket savolga hech kim javob bermasa, quiz avtomatik to‘xtaydi.\n\n"
        "Boshqaruvchi: /stop · /skip",
        reply_markup=group_home_keyboard(),
    )


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat = update.effective_chat

    # A single emoji may animate in Telegram clients; show it only once per
    # running process so /start does not become noisy.
    welcome_key = (chat.id, update.effective_user.id if update.effective_user else 0)
    if welcome_key not in WELCOME_SHOWN:
        WELCOME_SHOWN.add(welcome_key)
        try:
            await context.bot.send_message(chat_id=chat.id, text="👋")
        except Exception:
            pass

    if is_group_chat(chat):
        await send_group_home(update.message)
    else:
        await send_home(update.message)


async def home_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await send_home(query.message)


async def send_new_quiz_prompt(message):
    await message.reply_text(
        "📄 Yangi test\n\n"
        "PDF yoki Word (.docx) test faylini yuboring.\n\n"
        "✅ Asosiy taniladigan javob formatlari:\n"
        "• Javob: B / Answer: B\n"
        "• +B) ... / *B) ... / ✓B) ...\n\n"
        "Savol raqamida nuqta yoki bo‘sh joy tushib qolsa ham bot imkon qadar savolni taniydi. "
        "Bot to‘g‘ri javobni taxmin qilmaydi.\n\n"
        "🎁 Bepul: oyiga 1 ta yangi test import. Saqlangan testlarni cheksiz ishlash mumkin.",
        reply_markup=home_button(),
    )


async def new_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if is_group_chat(update.effective_chat):
        ok, host_name = await claim_group_host_with_menu(update.effective_chat.id, update.effective_user, context.bot)
        if not ok:
            await update.message.reply_text(
                f"🔒 Hozir {host_name} guruh quizini boshqaryapti."
            )
            return
        await update.message.reply_text(
            "📄 PDF yoki DOCX test faylini shu guruhga yuboring.\n"
            "Siz quiz boshqaruvchisisiz."
        )
    else:
        await send_new_quiz_prompt(update.message)


async def new_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await send_new_quiz_prompt(query.message)


# Keep the old callback name because other parts of the bot already use it.
async def help_upload_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await send_new_quiz_prompt(query.message)


async def show_saved_quizzes(message, tg_user, group_mode: bool = False):
    if not db.is_enabled():
        # Database is optional during the transition; keep old RAM behavior usable.
        session = GROUP_DATA.get(message.chat.id) if group_mode else USER_DATA.get(tg_user.id)
        if session:
            target = "g_current" if group_mode else "groups"
            await message.reply_text(
                "⚠️ Doimiy database hali ulanmagan. Hozirgi test faqat vaqtinchalik xotirada.\n\n"
                f"📄 {session.get('filename', 'Test')}\n"
                f"✅ {len(session.get('questions', []))} ta savol",
                reply_markup=InlineKeyboardMarkup(
                    [[InlineKeyboardButton("📚 Joriy test", callback_data=target)]]
                ),
            )
        else:
            await message.reply_text(
                "⚠️ Doimiy database hali ulanmagan. Avval yangi test yuklang."
            )
        return

    quizzes = await db.list_quizzes(tg_user.id, limit=20)
    if not quizzes:
        rows = []
        if group_mode:
            rows.append([InlineKeyboardButton("📄 Test yuklash", callback_data="g_new")])
            rows.append([InlineKeyboardButton("🏠 Guruh menyusi", callback_data="g_home")])
        else:
            rows.append([InlineKeyboardButton("📄 Yangi test", callback_data="menu_new")])
            rows.append([InlineKeyboardButton("🏠 Bosh menyu", callback_data="menu_home")])

        await message.reply_text(
            "📚 Sizda hali saqlangan test yo‘q.\n\n"
            "PDF/DOCX testni bir marta yuklasangiz, u keyingi safar shu yerda chiqadi.",
            reply_markup=InlineKeyboardMarkup(rows),
        )
        return

    rows = []
    prefix = "gload" if group_mode else "pquiz"
    for quiz in quizzes:
        name = quiz["name"]
        if len(name) > 32:
            name = name[:29] + "..."
        rows.append(
            [
                InlineKeyboardButton(
                    f"📘 {name} · {quiz['question_count']}",
                    callback_data=f"{prefix}:{quiz['id']}",
                )
            ]
        )

    rows.append(
        [
            InlineKeyboardButton(
                "🏠 Guruh menyusi" if group_mode else "🏠 Bosh menyu",
                callback_data="g_home" if group_mode else "menu_home",
            )
        ]
    )

    await message.reply_text(
        "📚 Saqlangan testlaringiz\n\nBoshlash uchun testni tanlang:",
        reply_markup=InlineKeyboardMarkup(rows),
    )


async def quizzes_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    group_mode = is_group_chat(update.effective_chat)
    if group_mode:
        ok, host_name = await claim_group_host_with_menu(update.effective_chat.id, update.effective_user, context.bot)
        if not ok:
            await update.message.reply_text(
                f"🔒 Hozir {host_name} guruh quizini boshqaryapti."
            )
            return
    await show_saved_quizzes(
        update.message,
        update.effective_user,
        group_mode=group_mode,
    )


async def quizzes_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await show_saved_quizzes(query.message, update.effective_user, group_mode=False)


async def group_saved_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    ok, host_name = await claim_group_host_with_menu(update.effective_chat.id, update.effective_user, context.bot)
    if not ok:
        await query.answer(f"Hozir {host_name} boshqaryapti.", show_alert=True)
        return
    await query.answer()
    await show_saved_quizzes(query.message, update.effective_user, group_mode=True)


async def private_quiz_detail_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    quiz_id = int(query.data.split(":")[1])
    quiz = await db.load_quiz(update.effective_user.id, quiz_id) if db.is_enabled() else None
    if not quiz:
        await query.message.reply_text("❌ Test topilmadi yoki sizga tegishli emas.")
        return

    pro = await user_has_pro(update.effective_user.id)
    rows = [[InlineKeyboardButton("▶️ Boshlash", callback_data=f"pload:{quiz_id}")]]
    if pro:
        rows += [
            [InlineKeyboardButton("🧠 Zaif savollar", callback_data=f"pweak:{quiz_id}"),
             InlineKeyboardButton("⭐ Belgilanganlar", callback_data=f"pbookmarks:{quiz_id}")],
            [InlineKeyboardButton("📊 PRO statistika", callback_data=f"presults:{quiz_id}")],
            [InlineKeyboardButton("📄 Nusxa yaratish", callback_data=f"pduplicate:{quiz_id}"),
             InlineKeyboardButton("♻️ Natijalarni tozalash", callback_data=f"presetstats:{quiz_id}")],
        ]
    else:
        rows.append([InlineKeyboardButton("👑 PRO imkoniyatlari", callback_data="pro_info")])
    rows += [
        [InlineKeyboardButton("✏️ Nomini o‘zgartirish", callback_data=f"prename:{quiz_id}")],
        [InlineKeyboardButton("🗑 O‘chirish", callback_data=f"pdelete:{quiz_id}")],
        [InlineKeyboardButton("← Testlarim", callback_data="menu_quizzes")],
    ]

    await query.message.reply_text(
        f"📘 {quiz['name']}\n\n❓ Savollar: {len(quiz['questions'])}\n📄 Manba: {quiz['source_filename']}",
        reply_markup=InlineKeyboardMarkup(rows),
    )


async def private_duplicate_quiz_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()
    if not db.is_enabled():
        await query.message.reply_text("❌ Database ulanmagan.")
        return
    quiz_id = int(query.data.split(":")[1])
    try:
        result = await db.duplicate_quiz(update.effective_user.id, quiz_id)
    except Exception:
        logging.exception("Quiz duplication failed")
        result = None
    if not result:
        await query.message.reply_text("❌ Test nusxasini yaratib bo‘lmadi.")
        return
    await query.message.reply_text(
        f"✅ Nusxa yaratildi: {result['name']}",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("📘 Nusxani ochish", callback_data=f"pquiz:{result['quiz_id']}")],
            [InlineKeyboardButton("📚 Testlarim", callback_data="menu_quizzes")],
        ]),
    )


async def private_reset_stats_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()
    quiz_id = int(query.data.split(":")[1])
    await query.message.reply_text(
        "♻️ Shu testning urinishlari va zaif-savol tarixini tozalaysizmi?\n\n"
        "⭐ Belgilangan savollar saqlanib qoladi.",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("♻️ Ha, tozalash", callback_data=f"presetstatsyes:{quiz_id}")],
            [InlineKeyboardButton("← Bekor qilish", callback_data=f"pquiz:{quiz_id}")],
        ]),
    )


async def private_reset_stats_yes_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    user_id = update.effective_user.id

    if not await user_has_pro(user_id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()

    quiz_id = int(query.data.split(":")[1])

    try:
        reset_info = await db.reset_quiz_progress(
            user_id,
            quiz_id,
            keep_bookmarks=True,
            verify=True,
        )
    except Exception:
        logging.exception("Reset quiz progress failed")
        reset_info = {"ok": False}

    if not reset_info.get("ok"):
        await query.message.reply_text(
            "❌ Natijalarni tozalab bo‘lmadi. Ma’lumotlar bazasi o‘zgarishlarni "
            "tasdiqlamadi."
        )
        return

    # Remove stale in-memory result/weak caches for the currently loaded quiz.
    session = USER_DATA.get(user_id)
    if session and session.get("saved_quiz_id") and int(session["saved_quiz_id"]) == quiz_id:
        session["results"] = {}
        session.pop("special_mode", None)

    context.user_data.pop(f"weak_positions:{quiz_id}", None)

    # Delete previously generated statistics summary messages from this chat.
    stats_key = f"stats_message_ids:{quiz_id}"
    old_message_ids = context.user_data.pop(stats_key, [])
    for message_id in old_message_ids:
        try:
            await context.bot.delete_message(
                chat_id=update.effective_chat.id,
                message_id=message_id,
            )
        except Exception:
            pass

    # Replace the confirmation card itself so there is no stale "are you sure?" UI.
    try:
        await query.edit_message_text(
            "✅ Natijalar tozalandi.\n\n"
            "Urinishlar: 0\n"
            "Zaif-savol tarixi: 0\n"
            "⭐ Belgilangan savollar saqlandi.\n\n"
            "Quyidagi tugma bazadan yangidan hisoblangan statistikani ochadi.",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("📊 Yangilangan statistika", callback_data=f"presults:{quiz_id}")],
                [InlineKeyboardButton("← Testga qaytish", callback_data=f"pquiz:{quiz_id}")],
            ]),
        )
    except Exception:
        await query.message.reply_text(
            "✅ Natijalar tozalandi. Urinishlar va zaif-savol tarixi 0 ga tushirildi.",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("📊 Yangilangan statistika", callback_data=f"presults:{quiz_id}")],
                [InlineKeyboardButton("← Testga qaytish", callback_data=f"pquiz:{quiz_id}")],
            ]),
        )


async def private_rename_quiz_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    context.user_data["rename_quiz_id"] = int(query.data.split(":")[1])
    await query.message.reply_text("✏️ Yangi nomni yozing. Bekor qilish uchun /cancel yuboring.")


async def cancel_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    pending = context.user_data.pop("rename_quiz_id", None)
    range_pending = context.user_data.pop("custom_range_request", None)
    random_pending = context.user_data.pop("custom_random_request", None)
    await update.message.reply_text(
        "✅ Amal bekor qilindi." if (pending or range_pending or random_pending) else "ℹ️ Bekor qilinadigan amal yo‘q."
    )


async def handle_plain_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if await apply_custom_range_text(update, context):
        return
    if await apply_custom_random_text(update, context):
        return

    quiz_id = context.user_data.get("rename_quiz_id")
    if not quiz_id:
        return
    name = (update.message.text or "").strip()
    if not name:
        await update.message.reply_text("❌ Nom bo‘sh bo‘lishi mumkin emas.")
        return
    if len(name) > 120:
        await update.message.reply_text("❌ Nom 120 belgidan qisqa bo‘lsin.")
        return
    ok = await db.rename_quiz(update.effective_user.id, int(quiz_id), name)
    if ok:
        context.user_data.pop("rename_quiz_id", None)
        await update.message.reply_text(
            f"✅ Test nomi o‘zgartirildi: {name}",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("📚 Testlarim", callback_data="menu_quizzes")]]),
        )
    else:
        await update.message.reply_text("❌ Test nomini o‘zgartirib bo‘lmadi.")


async def private_delete_quiz_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    quiz_id = int(query.data.split(":")[1])
    quiz = await db.load_quiz(update.effective_user.id, quiz_id) if db.is_enabled() else None
    if not quiz:
        await query.message.reply_text("❌ Test topilmadi.")
        return
    await query.message.reply_text(
        f"🗑 “{quiz['name']}” testini butunlay o‘chirasizmi?",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("🗑 Ha, o‘chirish", callback_data=f"pdeleteyes:{quiz_id}")],
            [InlineKeyboardButton("← Bekor qilish", callback_data=f"pquiz:{quiz_id}")],
        ]),
    )


async def private_delete_quiz_yes_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    quiz_id = int(query.data.split(":")[1])
    ok = await db.delete_quiz(update.effective_user.id, quiz_id)
    if ok:
        session = USER_DATA.get(update.effective_user.id)
        if session and session.get("saved_quiz_id") == quiz_id:
            USER_DATA.pop(update.effective_user.id, None)
        await query.message.reply_text("✅ Test o‘chirildi.", reply_markup=home_button())
    else:
        await query.message.reply_text("❌ Testni o‘chirib bo‘lmadi.")


async def private_quiz_results_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    quiz_id = int(query.data.split(":")[1])
    quiz = await db.load_quiz(update.effective_user.id, quiz_id) if db.is_enabled() else None
    if not quiz:
        await query.message.reply_text("❌ Test topilmadi.")
        return

    attempts = await db.list_quiz_attempts(update.effective_user.id, quiz_id, 10)
    pro = await user_has_pro(update.effective_user.id)
    lines = [f"📊 {quiz['name']}"]

    if pro:
        stats = await db.get_quiz_advanced_stats(update.effective_user.id, quiz_id)
        trend = await db.get_quiz_progress_trend(update.effective_user.id, quiz_id)
        delta = int(trend.get("improvement", 0) or 0)
        delta_text = f"+{delta}%" if delta > 0 else f"{delta}%"
        lines += [
            "",
            "👑 PRO statistika",
            f"Urinishlar: {stats.get('attempt_count', 0)}",
            f"O‘rtacha: {stats.get('average_percent', 0)}%",
            f"Eng yaxshi: {stats.get('best_percent', 0)}%",
            f"Oxirgi: {stats.get('latest_percent', 0)}%",
            f"📈 Dastlabki → oxirgi o‘rtacha: {trend.get('first_average', 0)}% → {trend.get('recent_average', 0)}% ({delta_text})",
            f"🔥 Eng uzun seriya: {stats.get('best_streak', 0)}",
            f"✅ Barqaror savollar: {stats.get('stable_questions', 0)}",
            f"🔴 Juda qiyin: {trend.get('red', 0)}",
            f"🟠 Qiyin: {trend.get('orange', 0)}",
            f"🟡 Takrorlash kerak: {trend.get('yellow', 0)}",
            f"⭐ Belgilangan: {stats.get('bookmarks', 0)}",
        ]

    if not attempts:
        lines += ["", "Hozircha natija yo‘q."]
    else:
        lines += ["", "Oxirgi natijalar:"]
        for row in attempts:
            mode = {"group": "👥 Guruh", "review": "🧠 Xatolar", "weak": "🧠 Zaif", "bookmarks": "⭐ Belgilangan"}.get(row.get("mode"), "📘 Shaxsiy")
            sec = row.get("section_index")
            sec_text = f" · {sec + 1}-bo‘lim" if sec is not None else ""
            lines.append(f"{mode}{sec_text}: {row['correct']}/{row['total']} ({row['percent']}%) · 🔥 {row['best_streak']}")

    rows = []
    if pro:
        rows += [[InlineKeyboardButton("🧠 Zaif savollar", callback_data=f"pweak:{quiz_id}"),
                  InlineKeyboardButton("⭐ Belgilanganlar", callback_data=f"pbookmarks:{quiz_id}")]]
    rows += [[InlineKeyboardButton("← Testga qaytish", callback_data=f"pquiz:{quiz_id}")],
             [InlineKeyboardButton("📚 Testlarim", callback_data="menu_quizzes")]]
    sent = await query.message.reply_text(
        "\n".join(lines),
        reply_markup=InlineKeyboardMarkup(rows),
    )
    key = f"stats_message_ids:{quiz_id}"
    ids = context.user_data.setdefault(key, [])
    ids.append(sent.message_id)
    # Keep the list bounded.
    context.user_data[key] = ids[-20:]


def _questions_by_positions(quiz: dict, positions: List[int]) -> List[dict]:
    wanted = set(int(x) for x in positions)
    return [q for q in quiz.get("questions", []) if int(q.get("position") or 0) in wanted]


async def _prepare_pro_study_set(update: Update, questions: List[dict], quiz: dict, mode_name: str, title: str):
    user_id = update.effective_user.id
    prefs = await effective_preferences(user_id)
    USER_DATA[user_id] = {
        "chat_id": update.effective_chat.id,
        "filename": quiz["source_filename"],
        "questions": questions,
        "custom_questions": questions,
        "custom_range_label": title,
        "warnings": [],
        "group_size": len(questions),
        "groups": [list(questions)],
        "active": None,
        "results": {},
        "saved_quiz_id": int(quiz["id"]),
        "owner_username": update.effective_user.username,
        "owner_full_name": update.effective_user.full_name,
        "shuffle_questions": bool(prefs.get("shuffle_questions")),
        "shuffle_options": bool(prefs.get("shuffle_options")),
        "quiz_mode": None,
        "parser_total_blocks": len(questions),
        "default_group_size": int(prefs.get("default_group_size", 50)),
        "default_timer": int(prefs.get("default_timer", 30)),
        "special_mode": mode_name,
    }


async def weak_questions_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()
    quiz_id = int(query.data.split(":")[1])
    quiz = await db.load_quiz(update.effective_user.id, quiz_id)
    if not quiz:
        await query.message.reply_text("❌ Test topilmadi.")
        return

    rows = await db.list_question_performance(update.effective_user.id, quiz_id)
    buckets = {"red": [], "orange": [], "yellow": []}
    for row in rows:
        bucket = pro_features.weak_bucket(
            row.get("attempts", 0), row.get("correct", 0), row.get("wrong", 0),
            row.get("unanswered", 0), row.get("last_result", ""),
        )
        if bucket:
            buckets[bucket].append(int(row["question_position"]))

    positions = buckets["red"] + buckets["orange"] + buckets["yellow"]
    if not positions:
        await query.message.reply_text(
            "🧠 Hozircha zaif savollar yo‘q.\n\n"
            "Testlarni ishlaganingiz sari bot savollar bo‘yicha natijalarni yig‘adi.",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("← Testga qaytish", callback_data=f"pquiz:{quiz_id}")]]),
        )
        return

    text = (
        "🧠 ZAIF SAVOLLAR\n\n"
        f"🔴 Juda qiyin: {len(buckets['red'])}\n"
        f"🟠 Qiyin: {len(buckets['orange'])}\n"
        f"🟡 Takrorlash kerak: {len(buckets['yellow'])}\n\n"
        f"Jami mashq uchun: {len(positions)} ta"
    )
    context.user_data[f"weak_positions:{quiz_id}"] = positions
    await query.message.reply_text(
        text,
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("▶️ Zaif savollarni mashq qilish", callback_data=f"pweakstart:{quiz_id}")],
            [InlineKeyboardButton("← Testga qaytish", callback_data=f"pquiz:{quiz_id}")],
        ]),
    )


async def weak_start_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()
    quiz_id = int(query.data.split(":")[1])
    quiz = await db.load_quiz(update.effective_user.id, quiz_id)
    positions = context.user_data.get(f"weak_positions:{quiz_id}")
    if not positions:
        rows = await db.list_question_performance(update.effective_user.id, quiz_id)
        positions = [int(r["question_position"]) for r in rows if pro_features.weak_bucket(
            r.get("attempts",0), r.get("correct",0), r.get("wrong",0), r.get("unanswered",0), r.get("last_result","")
        )]
    questions = _questions_by_positions(quiz, positions or []) if quiz else []
    if not questions:
        await query.message.reply_text("✅ Mashq qilish uchun zaif savol qolmagan.")
        return
    await _prepare_pro_study_set(update, questions, quiz, "weak", "Zaif savollar")
    await query.message.reply_text(
        f"🧠 {len(questions)} ta zaif savol tayyor.\n\nDavom eting:",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("⏱ Taymer tanlash", callback_data="group:0")]]),
    )


async def bookmarks_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()

    quiz_id = int(query.data.split(":")[1])
    quiz = await db.load_quiz(update.effective_user.id, quiz_id)
    if not quiz:
        await query.message.reply_text("❌ Test topilmadi.")
        return

    positions = await db.list_bookmark_positions(update.effective_user.id, quiz_id)
    if not positions:
        await query.message.reply_text(
            "⭐ Hozircha belgilangan savol yo‘q.\n\n"
            "PRO test ishlayotganda savol ostidagi ⭐ Keyin ko‘rish tugmasini bosing. "
            "Tugma savolga javob berganingizdan keyin ham qoladi.",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("← Testga qaytish", callback_data=f"pquiz:{quiz_id}")]
            ]),
        )
        return

    text = bookmark_preview_text(quiz, positions)
    await query.message.reply_text(
        text,
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton(
                f"▶️ Barcha {len(positions)} ta belgilangan savolni mashq qilish",
                callback_data=f"pbookstart:{quiz_id}",
            )],
            [InlineKeyboardButton("← Testga qaytish", callback_data=f"pquiz:{quiz_id}")],
        ]),
    )


async def bookmarks_start_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()
    quiz_id = int(query.data.split(":")[1])
    quiz = await db.load_quiz(update.effective_user.id, quiz_id)
    positions = await db.list_bookmark_positions(update.effective_user.id, quiz_id) if quiz else []
    questions = _questions_by_positions(quiz, positions) if quiz else []
    if not questions:
        await query.message.reply_text("⭐ Belgilangan savol topilmadi.")
        return
    await _prepare_pro_study_set(update, questions, quiz, "bookmarks", "Belgilangan savollar")
    await query.message.reply_text(
        f"⭐ {len(questions)} ta belgilangan savol tayyor.\n\nDavom eting:",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("⏱ Taymer tanlash", callback_data="group:0")]]),
    )


async def private_load_saved_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if not db.is_enabled():
        await query.message.reply_text("❌ Database ulanmagan.")
        return

    quiz_id = int(query.data.split(":")[1])
    quiz = await db.load_quiz(update.effective_user.id, quiz_id)
    if not quiz:
        await query.message.reply_text("❌ Test topilmadi yoki sizga tegishli emas.")
        return

    prefs = await effective_preferences(update.effective_user.id)
    USER_DATA[update.effective_user.id] = {
        "chat_id": update.effective_chat.id,
        "filename": quiz["source_filename"],
        "questions": quiz["questions"],
        "warnings": [],
        "group_size": None,
        "groups": [],
        "active": None,
        "results": {},
        "saved_quiz_id": quiz_id,
        "owner_username": update.effective_user.username,
        "owner_full_name": update.effective_user.full_name,
        "shuffle_questions": bool(prefs.get("shuffle_questions")),
        "shuffle_options": bool(prefs.get("shuffle_options")),
        "quiz_mode": None,
        "parser_total_blocks": len(quiz["questions"]),
        "default_group_size": int(prefs.get("default_group_size", 50)),
        "default_timer": int(prefs.get("default_timer", 30)),
    }

    await query.message.reply_text(
        f"📘 {quiz['name']}\n"
        f"✅ {len(quiz['questions'])} ta savol yuklandi.\n\n"
        "Har bir guruhda nechta savol bo‘lsin?",
        reply_markup=group_size_keyboard(group_mode=False, default_size=USER_DATA[update.effective_user.id].get("default_group_size")),
    )


async def group_load_saved_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = update.effective_chat.id
    ok, host_name = await claim_group_host_with_menu(chat_id, update.effective_user, context.bot)
    if not ok:
        await query.answer(
            f"Bu guruh quizini hozir {host_name} boshqaryapti.",
            show_alert=True,
        )
        return
    await query.answer()

    if not db.is_enabled():
        await query.message.reply_text("❌ Database ulanmagan.")
        return

    quiz_id = int(query.data.split(":")[1])
    quiz = await db.load_quiz(update.effective_user.id, quiz_id)
    if not quiz:
        await query.message.reply_text("❌ Test topilmadi yoki sizga tegishli emas.")
        return

    prefs = await effective_preferences(update.effective_user.id)
    GROUP_DATA[chat_id] = {
        "chat_id": chat_id,
        "filename": quiz["source_filename"],
        "questions": quiz["questions"],
        "warnings": [],
        "group_size": None,
        "groups": [],
        "active": None,
        "results": {},
        "controller_id": update.effective_user.id,
        "controller_name": update.effective_user.full_name,
        "host_last_activity": time.time(),
        "last_leaderboard_text": None,
        "saved_quiz_id": quiz_id,
        "shuffle_questions": bool(prefs.get("shuffle_questions")),
        "shuffle_options": bool(prefs.get("shuffle_options")),
        "quiz_mode": None,
        "parser_total_blocks": len(quiz["questions"]),
        "default_group_size": int(prefs.get("default_group_size", 50)),
        "default_timer": int(prefs.get("default_timer", 30)),
    }

    await query.message.reply_text(
        f"✅ {update.effective_user.full_name} boshqaruvchi bo‘ldi.\n\n"
        f"📘 {quiz['name']}\n"
        f"❓ {len(quiz['questions'])} ta savol\n\n"
        "Har bir bo‘limda nechta savol bo‘lsin?",
        reply_markup=group_size_keyboard(group_mode=True, default_size=GROUP_DATA[chat_id].get("default_group_size")),
    )


async def show_continue(message, user_id: int):
    session = USER_DATA.get(user_id)
    if not session:
        await message.reply_text(
            "▶️ Davom ettirish uchun faol test topilmadi.",
            reply_markup=InlineKeyboardMarkup(
                [
                    [InlineKeyboardButton("📄 Yangi test", callback_data="menu_new")],
                    [InlineKeyboardButton("🏠 Bosh menyu", callback_data="menu_home")],
                ]
            ),
        )
        return

    if session.get("active"):
        active = session["active"]
        if active.get("paused"):
            await message.reply_text(
                f"⏸ Quiz pauzada.\n\n"
                f"Guruh: {active['group_index'] + 1}\n"
                f"Keyingi savol: {active['current'] + 1}/{len(active['questions'])}\n\n"
                "Davom ettirish uchun /resume yuboring.",
                reply_markup=home_button(),
            )
        else:
            await message.reply_text(
                f"▶️ Quiz hozir davom etmoqda.\n\n"
                f"Guruh: {active['group_index'] + 1}\n"
                f"Savol: {active['current'] + 1}/{len(active['questions'])}\n\n"
                "Joriy Telegram polliga javob bering yoki vaqt tugashini kuting.",
                reply_markup=home_button(),
            )
        return

    if session.get("groups"):
        await show_groups(message, user_id)
        return

    await message.reply_text(
        "▶️ Test yuklangan, lekin guruh hajmi hali tanlanmagan.\n\n"
        "Har bir guruhda nechta savol bo‘lsin?",
        reply_markup=group_size_keyboard(default_size=session.get("default_group_size")),
    )


async def continue_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_continue(update.message, update.effective_user.id)


async def continue_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await show_continue(query.message, update.effective_user.id)


async def show_progress(message, user_id: int):
    if db.is_enabled():
        try:
            attempts = await db.list_recent_attempts(user_id, limit=10)
        except Exception:
            logging.exception("Could not load attempt history")
            attempts = []

        pro = await user_has_pro(user_id)
        lines = []
        if pro:
            try:
                stats = await db.get_user_advanced_stats(user_id)
            except Exception:
                logging.exception("Could not load Pro stats")
                stats = {}
            lines += [
                "📊 PRO NATIJALAR",
                "",
                f"Urinishlar: {stats.get('attempt_count', 0)}",
                f"Ishlangan testlar: {stats.get('quiz_count', 0)}",
                f"O‘rtacha: {stats.get('average_percent', 0)}%",
                f"Eng yaxshi: {stats.get('best_percent', 0)}%",
                f"🔥 Eng uzun seriya: {stats.get('best_streak', 0)}",
                f"⚠️ Zaif savollar: {stats.get('weak_questions', 0)}",
                f"⭐ Belgilangan: {stats.get('bookmarks', 0)}",
            ]

        if attempts:
            lines += (["", "Oxirgi natijalar:"] if lines else ["📊 Oxirgi natijalar", ""])
            for row in attempts:
                name = row.get("quiz_name") or "Test"
                mode = {"group": "👥", "review": "🧠", "weak": "🧠", "bookmarks": "⭐"}.get(row.get("mode"), "📘")
                section = row.get("section_index")
                section_text = f" · {section + 1}-bo‘lim" if section is not None else ""
                lines.append(f"{mode} {name}{section_text} — {row['correct']}/{row['total']} ({row['percent']}%)")

        if lines:
            await message.reply_text("\n".join(lines), reply_markup=home_button())
            return

    session = USER_DATA.get(user_id)
    results = (session or {}).get("results", {})
    if not results:
        await message.reply_text("📊 Hozircha natijalar yo‘q.", reply_markup=home_button())
        return

    lines = ["📊 Joriy sessiya natijalari\n"]
    for idx in sorted(results):
        result = results[idx]
        unanswered = result.get("unanswered", 0)
        lines.append(
            f"📘 {idx + 1}-guruh — {result['correct']}/{result['total']} "
            f"({result['percent']}%)" + (f" · javobsiz {unanswered}" if unanswered else "")
        )
    await message.reply_text("\n".join(lines), reply_markup=home_button())


async def progress_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await show_progress(update.message, update.effective_user.id)


async def progress_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await show_progress(query.message, update.effective_user.id)


async def send_group_mode_info(message, context: ContextTypes.DEFAULT_TYPE):
    chat = message.chat

    if is_group_chat(chat):
        await send_group_home(message)
        return

    me = await context.bot.get_me()
    add_url = f"https://t.me/{me.username}?startgroup=quiz" if me.username else None

    rows = []
    if add_url:
        rows.append([InlineKeyboardButton("➕ Botni guruhga qo‘shish", url=add_url)])
    rows.append([InlineKeyboardButton("🏠 Bosh menyu", callback_data="menu_home")])

    await message.reply_text(
        "👥 Guruh testi\n\n"
        "Botni Telegram guruhiga qo‘shing va guruh ichida /group yoki /start yuboring.\n"
        "So‘ng PDF/DOCX test faylini guruhga yuboring.\n\n"
        "Guruh rejimida:\n"
        "• hamma bir xil savolni ko‘radi;\n"
        "• javob bergan odam keyingi savolni erta boshlatmaydi;\n"
        "• savol tanlangan vaqtning oxirigacha ochiq turadi;\n"
        "• yakunda individual reyting chiqadi.",
        reply_markup=InlineKeyboardMarkup(rows),
    )


async def group_mode_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if is_group_chat(update.effective_chat):
        ok, host_name = await claim_group_host_with_menu(update.effective_chat.id, update.effective_user, context.bot)
        if not ok:
            await update.message.reply_text(
                f"🔒 Bu guruh quizini hozir {host_name} boshqaryapti. "
                "Quiz tugagach yoki sessiya bo‘shatilgach yana urinib ko‘ring."
            )
            return
        await update.message.reply_text(
            f"🎮 {update.effective_user.full_name} guruh quiz boshqaruvchisi bo‘ldi."
        )
        await send_group_home(update.message)
        return
    await send_group_mode_info(update.message, context)


async def group_mode_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await send_group_mode_info(query.message, context)


async def group_home_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await send_group_home(query.message)


async def group_new_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    ok, host_name = await claim_group_host_with_menu(update.effective_chat.id, update.effective_user, context.bot)
    if not ok:
        await query.answer(f"Hozir {host_name} boshqaryapti.", show_alert=True)
        return
    await query.answer()
    await query.message.reply_text(
        "📄 PDF yoki DOCX test faylini shu guruhga yuboring.\n\n"
        "Siz quiz boshqaruvchisisiz. Boshqa foydalanuvchi sizning setupingizni almashtira olmaydi."
    )


async def group_current_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    chat_id = update.effective_chat.id
    session = GROUP_DATA.get(chat_id)

    if not session:
        await query.message.reply_text(
            "📚 Bu guruhda hali test yuklanmagan.",
            reply_markup=group_home_keyboard(),
        )
        return

    total = len(session.get("questions", []))
    warnings = len(session.get("warnings", []))
    text = (
        f"📚 Joriy guruh testi\n\n"
        f"📄 {session.get('filename', 'Test')}\n"
        f"✅ Quizga tayyor: {total}\n"
        f"⚠️ Muammoli: {warnings}"
    )

    if session.get("groups"):
        text += f"\n📦 Guruhlar: {len(session['groups'])}"

    await query.message.reply_text(text, reply_markup=group_home_keyboard())


async def group_leaderboard_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    session = GROUP_DATA.get(update.effective_chat.id)

    if not session or not session.get("last_leaderboard_text"):
        await query.message.reply_text(
            "🏆 Hozircha yakunlangan guruh testi yo‘q.",
            reply_markup=group_home_keyboard(),
        )
        return

    await query.message.reply_text(
        session["last_leaderboard_text"],
        reply_markup=group_home_keyboard(),
    )


async def group_help_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "❓ Guruh rejimi\n\n"
        "1) PDF/DOCX fayl yuboring.\n"
        "2) 30 / 40 / 50 / 100 yoki PRO custom range/random count ni tanlang.\n"
        "3) Quiz bo‘limini tanlang.\n"
        "4) Vaqtni tanlang.\n"
        "5) ▶️ START ni bosing.\n\n"
        "Savol vaqt tugaguncha ochiq turadi. Hamma shu vaqt ichida javob beradi. "
        "Keyingi savol avtomatik ravishda vaqt tugagach chiqadi.\n\n"
        "⏸ /pause — quizni vaqtincha to‘xtatish\n"
        "▶️ /resume — pauzadagi quizni davom ettirish\n"
        "🛑 /stop — quizni butunlay tugatish\n"
        "⏭ /skip — joriy savolni o‘tkazish\n"
        f"😴 {GROUP_EMPTY_STOP_THRESHOLD} ta ketma-ket savolga hech kim javob bermasa, bot o‘zi to‘xtaydi.",
        reply_markup=group_home_keyboard(),
    )


async def settings_keyboard(user_id: int) -> InlineKeyboardMarkup:
    prefs = await effective_preferences(user_id)
    pro = await user_has_pro(user_id)
    rows = [
        [InlineKeyboardButton(
            f"{'✅' if prefs['shuffle_questions'] else '⬜'} Savollarni aralashtirish",
            callback_data="prefs_qshuffle",
        )],
        [InlineKeyboardButton(
            f"{'✅' if prefs['shuffle_options'] else '⬜'} Variantlarni aralashtirish",
            callback_data="prefs_oshuffle",
        )],
    ]
    if pro:
        rows += [
            [InlineKeyboardButton(
                f"👑 Standart bo‘lim: {prefs['default_group_size']}",
                callback_data="prefs_size",
            )],
            [InlineKeyboardButton(
                f"👑 Standart taymer: {format_duration(int(prefs['default_timer']))}",
                callback_data="prefs_timer",
            )],
        ]
    else:
        rows.append([InlineKeyboardButton("👑 PRO sozlamalari", callback_data="pro_info")])
    rows.append([InlineKeyboardButton("🏠 Bosh menyu", callback_data="menu_home")])
    return InlineKeyboardMarkup(rows)


async def send_settings(message, user_id: int):
    prefs = await effective_preferences(user_id)
    pro = await user_has_pro(user_id)
    pro_text = (
        f"\n👑 Standart bo‘lim: {prefs['default_group_size']}\n"
        f"👑 Standart taymer: {format_duration(int(prefs['default_timer']))}\n"
        if pro else
        "\n👑 PRO da standart bo‘lim va taymerni saqlash mumkin.\n"
    )
    await message.reply_text(
        "⚙️ Standart sozlamalar\n\n"
        "Aralashtirish sozlamalari yangi testlarga avtomatik qo‘llanadi. "
        "Mashq/Imtihon rejimi esa har bir test boshlanishidan oldin majburiy tanlanadi.\n\n"
        f"🔀 Savollar: {'yoqilgan' if prefs['shuffle_questions'] else 'o‘chirilgan'}\n"
        f"🔀 Variantlar: {'yoqilgan' if prefs['shuffle_options'] else 'o‘chirilgan'}"
        f"{pro_text}",
        reply_markup=await settings_keyboard(user_id),
    )


async def settings_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await send_settings(update.message, update.effective_user.id)


async def settings_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await send_settings(query.message, update.effective_user.id)


async def prefs_toggle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    prefs = await effective_preferences(update.effective_user.id)
    if query.data == "prefs_qshuffle":
        changes = {"shuffle_questions": not prefs["shuffle_questions"]}
    elif query.data == "prefs_oshuffle":
        changes = {"shuffle_options": not prefs["shuffle_options"]}
    else:
        return
    await db.update_user_preferences(
        update.effective_user.id,
        update.effective_user.username,
        update.effective_user.full_name,
        **changes,
    )
    await send_settings(query.message, update.effective_user.id)


async def prefs_size_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()
    await query.message.reply_text(
        "👑 Standart bo‘lim hajmini tanlang:",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("30", callback_data="prefs_set_size:30"),
             InlineKeyboardButton("40", callback_data="prefs_set_size:40")],
            [InlineKeyboardButton("50", callback_data="prefs_set_size:50"),
             InlineKeyboardButton("100", callback_data="prefs_set_size:100")],
        ]),
    )


async def prefs_set_size_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()
    value = int(query.data.split(":")[1])
    await db.update_user_preferences(
        update.effective_user.id,
        update.effective_user.username,
        update.effective_user.full_name,
        default_group_size=value,
    )
    await query.message.reply_text(f"✅ Standart bo‘lim: {value} ta savol.")
    await send_settings(query.message, update.effective_user.id)


async def prefs_timer_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()
    rows = [
        [InlineKeyboardButton("10 soniya", callback_data="prefs_set_timer:10"),
         InlineKeyboardButton("15 soniya", callback_data="prefs_set_timer:15")],
        [InlineKeyboardButton("20 soniya", callback_data="prefs_set_timer:20"),
         InlineKeyboardButton("30 soniya", callback_data="prefs_set_timer:30")],
        [InlineKeyboardButton("40 soniya", callback_data="prefs_set_timer:40"),
         InlineKeyboardButton("1 daqiqa", callback_data="prefs_set_timer:60")],
        [InlineKeyboardButton("2 daqiqa", callback_data="prefs_set_timer:120")],
    ]
    await query.message.reply_text("👑 Standart taymerni tanlang:", reply_markup=InlineKeyboardMarkup(rows))


async def prefs_set_timer_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    if not await user_has_pro(update.effective_user.id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    await query.answer()
    value = int(query.data.split(":")[1])
    await db.update_user_preferences(
        update.effective_user.id,
        update.effective_user.username,
        update.effective_user.full_name,
        default_timer=value,
    )
    await query.message.reply_text(f"✅ Standart taymer: {format_duration(value)}.")
    await send_settings(query.message, update.effective_user.id)


async def help_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("✅ Qo‘llab-quvvatlanadigan formatlar", callback_data="help_formats")],
        [InlineKeyboardButton("🤖 AI bilan formatlash", callback_data="help_ai")],
        [InlineKeyboardButton("🧪 Parser qanday ishlaydi?", callback_data="help_parser_info")],
        [InlineKeyboardButton("ℹ️ Bot haqida", callback_data="help_about"), InlineKeyboardButton("🔐 Maxfiylik", callback_data="help_privacy")],
        [InlineKeyboardButton("🏠 Bosh menyu", callback_data="menu_home")],
    ])


async def send_help(message):
    await message.reply_text(
        "❓ Yordam\n\n"
        "Testchi tayyor PDF/DOCX testlarni Telegram quizga aylantiradi, saqlaydi "
        "va shaxsiy yoki guruh rejimida ishlashga yordam beradi.\n\n"
        "⚠️ Har qanday hujjat avtomatik tanilmaydi. To‘g‘ri javob manbada aniq ko‘rsatilgan "
        "bo‘lishi kerak. Bot javobni taxmin qilmaydi.",
        reply_markup=await help_menu(),
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await send_help(update.message)


async def help_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await send_help(query.message)


async def help_formats_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "✅ QO‘LLAB-QUVVATLANADIGAN FORMATLAR\n\n"
        "Asosiy ko‘rinish:\n"
        "1. Savol matni\nA) Variant\nB) Variant\nC) Variant\nD) Variant\nJavob: B\n\n"
        "Parser quyidagilarga ham moslashadi:\n"
        "• 1. / 1) / №1 / Savol 1 / Question 1 / Вопрос 1\n"
        "• A) / A. / A: / A-\n"
        "• 2 tadan 10 tagacha variant\n"
        "• ko‘p qatorli savol va variantlar\n"
        "• bir qatorda A/B/C/D variantlar\n"
        "• Javob / Answer / Ответ\n"
        "• hujjat oxiridagi Javoblar / Answer key / Ответы\n"
        "• +B), *B), ✓B), ✔B) kabi aniq markerlar\n"
        "• Word jadvallari\n"
        "• real PDF highlight bilan belgilangan javobli jadvallar\n\n"
        "👑 PRO da noodatiy matnli formatlar va skaner PDFlar uchun AI fallback ham bor.\n\n"
        "⚠️ To‘g‘ri javob manbada aniq bo‘lmasa, bot uni taxmin qilmaydi.",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("← Yordam", callback_data="menu_help")]]),
    )


async def help_ai_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "🤖 AI PARSER\n\n"
        "Bepul tarifda avval standart parser ishlaydi. Noodatiy formatni tashqi AI yordamida moslashtirish mumkin.\n\n"
        "👑 PRO da botning o‘zi avtomatik ishlaydi:\n"
        "• oddiy parser birinchi urinadi;\n"
        "• muammoli matnli bloklar Gemini AI ga yuboriladi;\n"
        "• skaner/rasm PDF bo‘lsa, PRO AI sahifalarni o‘qishga urinadi;\n"
        "• AI javobni yechishi taqiqlangan; faqat manbada belgilangan javob qabul qilinadi.\n\n"
        f"Hozirgi PRO AI limiti: oyiga {PRO_AI_IMPORT_LIMIT} ta AI import.",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("← Yordam", callback_data="menu_help")]]),
    )


async def help_parser_info_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "🧪 PARSER QANDAY ISHLAYDI?\n\n"
        "1) Bot avval oddiy parser bilan savol bloklarini topadi.\n"
        "2) Savol matni va 2–10 ta variantni ajratadi.\n"
        "3) Faqat manbada aniq ko‘rsatilgan to‘g‘ri javobni qabul qiladi.\n"
        "4) Javoblar hujjat oxiridagi kalitda bo‘lsa, ularni moslashtirishga urinadi.\n"
        "5) Word jadvallari va haqiqiy PDF highlight belgilari ham tekshiriladi.\n"
        "6) Noaniq savollar muammoli deb hisobotda ko‘rsatiladi.\n"
        "7) 👑 PRO da oddiy parser yetarli bo‘lmasa AI fallback ishlaydi.\n\n"
        "⚠️ Bot to‘g‘ri javobni taxmin qilmaydi.\n\n"
        "/parser — joriy yuklangan test bo‘yicha parser hisobotini ko‘rish.",
        reply_markup=InlineKeyboardMarkup(
            [[InlineKeyboardButton("← Yordam", callback_data="menu_help")]]
        ),
    )


async def pro_info_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "👑 PRO IMKONIYATLARI\n\n"
        "♾ Oddiy matnli PDF/DOCX importlari cheklanmagan\n"
        f"🤖 Oyiga {PRO_AI_IMPORT_LIMIT} ta AI yordamidagi import\n"
        "🖼 Skaner/rasm PDFlarni AI orqali o‘qish\n"
        "🎯 Istalgan savol oralig‘i (masalan 121–170)\n"
        "🎲 Tasodifiy N ta savol\n"
        "⭐ Qiyin savollarni belgilash va alohida mashq qilish\n"
        "🧠 Uzoq muddatli zaif-savol mashqi\n"
        "📊 Kengaytirilgan statistika va progress trendi\n"
        "⚙️ Standart taymer/bo‘lim/aralashtirish sozlamalari\n"
        "📄 Saqlangan testdan nusxa yaratish va natijalarni reset qilish\n"
        "👥 Guruhda custom range/random count va host boshqaruvi\n\n"
        "💳 To‘lovlar hozircha yoqilmagan — avval barcha PRO funksiyalar test qilinadi.",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🏠 Bosh menyu", callback_data="menu_home")]]),
    )


async def help_about_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "ℹ️ TESTCHI HAQIDA\n\n"
        "PDF/DOCX testlarni Telegram quizga aylantiring, saqlang va istalgancha mashq qiling.\n\n"
        "⏱ Taymerlar\n🔀 Savol va variantlarni aralashtirish\n📖 Mashq / 📝 Imtihon rejimi\n"
        "❌ Xatolarni kamaytirib qayta mashq qilish\n⏸ Pauza / ▶️ davom / 🛑 to‘xtatish\n"
        "📊 Natijalar tarixi\n👥 Guruh testlari va reyting\n\n"
        "🎁 Bepul tarif: oyiga 1 ta yangi test importi. Saqlangan testlarni qayta ishlash cheklanmagan.",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("← Yordam", callback_data="menu_help")]]),
    )


async def help_privacy_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(
        "🔐 MAXFIYLIK\n\n"
        "Bot ishlashi uchun Telegram ID, Telegramdagi ism/username, siz saqlagan testlar va natijalar ma’lumotlar bazasida saqlanishi mumkin.\n\n"
        "Telefon raqamingiz hozirgi versiyada talab qilinmaydi va avtomatik olinmaydi.\n\n"
        "Original PDF/DOCX faylni doimiy saqlash shart emas; quiz uchun ajratilgan savol va variantlar saqlanadi.",
        reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("← Yordam", callback_data="menu_help")]]),
    )


def parser_report_text(session: Optional[dict]) -> str:
    if not session or not session.get("questions"):
        return "🧪 Hozircha parser hisoboti yo‘q. Avval test yuklang."
    total = int(session.get("parser_total_blocks") or (len(session.get("questions", [])) + len(session.get("warnings", []))))
    ready = len(session.get("questions", []))
    warnings = session.get("warnings", [])
    lines = [
        "🧪 Parser hisoboti",
        "",
        f"📄 {session.get('filename') or 'Test'}",
        f"🔧 Usul: {session.get('parser_method') or 'standart parser'}",
        f"🧩 Savol bloklari: {total}",
        f"✅ Quizga tayyor: {ready}",
        f"⚠️ Muammoli: {len(warnings)}",
    ]
    if warnings:
        lines.append("")
        lines.append("Muammolar:")
        lines.extend(f"• {w}" for w in warnings[:15])
        if len(warnings) > 15:
            lines.append(f"• ... yana {len(warnings) - 15} ta")
    return "\n".join(lines)


async def parser_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if is_group_chat(update.effective_chat):
        session = GROUP_DATA.get(update.effective_chat.id)
    else:
        session = USER_DATA.get(update.effective_user.id)
    await update.message.reply_text(parser_report_text(session))


async def group_parser_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await query.message.reply_text(parser_report_text(GROUP_DATA.get(update.effective_chat.id)))


async def plan_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    pro = await user_has_pro(user.id)

    if pro:
        usage = {"imports_used": 0, "recovered_questions": 0}
        if db.is_enabled() and not is_owner(user.id):
            try:
                usage = await db.get_ai_usage(user.id)
            except Exception:
                logging.exception("AI usage o‘qilmadi")
        owner_line = "🛠 OWNER / PRO TEST\n\n" if is_owner(user.id) else "👑 PRO\n\n"
        await update.message.reply_text(
            owner_line
            + "♾ Oddiy matnli PDF/DOCX importlari: cheklanmagan\n"
            + f"🤖 AI importlari: {usage['imports_used']}/{PRO_AI_IMPORT_LIMIT}\n"
            + "🎯 Maxsus savol oralig‘i: ochiq\n"
            + "🎲 Tasodifiy N ta savol: ochiq\n"
            + "⭐ Belgilangan savollar: ochiq\n"
            + "🧠 Zaif savollar mashqi: ochiq\n"
            + "📊 Kengaytirilgan statistika: ochiq\n"
            + "⚙️ Standart bo‘lim/taymer: ochiq\n\n"
            + "🖼 Skaner/rasm PDF: PRO AI orqali qo‘llab-quvvatlanadi (sahifa limiti mavjud)."
        )
        return

    if not db.is_enabled():
        await update.message.reply_text("⚠️ Database ulanmagan, tarif holatini aniqlab bo‘lmaydi.")
        return

    try:
        await db.ensure_user(user.id, user.username, user.full_name)
        status = await db.get_plan_status(user.id, free_limit=FREE_IMPORT_LIMIT)
    except Exception:
        logging.exception("Could not load plan")
        await update.message.reply_text("❌ Tarif holatini yuklab bo‘lmadi.")
        return

    await update.message.reply_text(
        "🎁 BEPUL\n\n"
        f"Bu oy yangi import: {status['imports_used']}/{FREE_IMPORT_LIMIT}\n"
        f"Qolgan: {status['imports_remaining']}\n\n"
        "✅ Saqlangan testlarni ishlash cheklanmagan.\n"
        "✅ Mashq/Imtihon, taymer, aralashtirish va guruh testlari ishlaydi.\n\n"
        "👑 PRO: cheksiz normal import, matnli va skan PDF uchun AI yordam, maxsus oraliq, "
        "tasodifiy N ta savol, bookmark, zaif savollar mashqi va kengaytirilgan statistika.\n\n"
        "💳 To‘lovlar hali yoqilmagan."
    )


async def grantpro_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_owner(update.effective_user.id):
        await update.message.reply_text("⛔ Bu buyruq faqat bot egasi uchun.")
        return

    if not db.is_enabled():
        await update.message.reply_text("❌ Database ulanmagan.")
        return

    if len(context.args) < 1:
        await update.message.reply_text(
            "Foydalanish:\n/grantpro TELEGRAM_ID KUN\n\n"
            "Misol: /grantpro 123456789 30"
        )
        return

    try:
        target_id = int(context.args[0])
        days = int(context.args[1]) if len(context.args) > 1 else 30
        if days < 1 or days > 3650:
            raise ValueError
    except ValueError:
        await update.message.reply_text("❌ Telegram ID yoki kun noto‘g‘ri.")
        return

    try:
        await db.grant_pro(target_id, days)
    except Exception:
        logging.exception("grantpro failed")
        await update.message.reply_text("❌ PRO faollashtirib bo‘lmadi.")
        return

    await update.message.reply_text(
        f"✅ {target_id} uchun PRO {days} kunga faollashtirildi."
    )


async def revokepro_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_owner(update.effective_user.id):
        await update.message.reply_text("⛔ Bu buyruq faqat bot egasi uchun.")
        return

    if not db.is_enabled():
        await update.message.reply_text("❌ Database ulanmagan.")
        return

    if not context.args:
        await update.message.reply_text(
            "Foydalanish:\n/revokepro TELEGRAM_ID"
        )
        return

    try:
        target_id = int(context.args[0])
        await db.revoke_pro(target_id)
    except Exception:
        logging.exception("revokepro failed")
        await update.message.reply_text("❌ PRO bekor qilinmadi.")
        return

    await update.message.reply_text(
        f"✅ {target_id} foydalanuvchi BEPUL tarifga qaytarildi."
    )


async def myid_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(f"Sizning Telegram ID: {update.effective_user.id}")


async def release_group_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_group_chat(update.effective_chat):
        await update.message.reply_text("/release faqat guruhda ishlaydi.")
        return

    chat_id = update.effective_chat.id
    session = GROUP_DATA.get(chat_id)
    if not session:
        await update.message.reply_text("🔓 Hozir guruh quiz boshqaruvchisi yo‘q.")
        return

    if update.effective_user.id != session.get("controller_id"):
        await update.message.reply_text("⛔ Boshqaruvni faqat joriy quiz egasi bo‘shata oladi.")
        return

    if session.get("active"):
        await update.message.reply_text("⚠️ Avval /stop bilan faol quizni to‘xtating.")
        return

    old_host_id = int(session["controller_id"])
    GROUP_DATA.pop(chat_id, None)
    await clear_host_command_menu(context.bot, chat_id, old_host_id)
    await update.message.reply_text(
        "🔓 Guruh quiz boshqaruvi bo‘shatildi.\n"
        "Endi boshqa foydalanuvchi /group orqali quiz boshlashi mumkin."
    )


async def group_release_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    session = GROUP_DATA.get(update.effective_chat.id)
    if not session:
        await query.answer("Sessiya yo‘q", show_alert=True)
        return
    if not group_controller_ok(update, session):
        await query.answer("Faqat boshqaruvchi sessiyani bo‘shata oladi.", show_alert=True)
        return
    if session.get("active"):
        await query.answer("Avval /stop bilan faol quizni to‘xtating.", show_alert=True)
        return
    await query.answer()
    chat_id = update.effective_chat.id
    old_host_id = int(session["controller_id"])
    GROUP_DATA.pop(chat_id, None)
    await clear_host_command_menu(context.bot, chat_id, old_host_id)
    await query.message.reply_text("🔓 Guruh quiz boshqaruvi bo‘shatildi. Endi boshqa foydalanuvchi boshlashi mumkin.")


def _warning_number(warning: str) -> Optional[int]:
    match = re.search(r"Question\s+(\d+)\s*:", warning or "", re.IGNORECASE)
    return int(match.group(1)) if match else None


def _remove_recovered_warnings(warnings: List[str], recovered: List[dict]) -> List[str]:
    recovered_numbers = {
        int(q["number"])
        for q in recovered
        if q.get("number") is not None
    }
    if not recovered_numbers:
        return list(warnings)

    result = []
    for warning in warnings:
        number = _warning_number(warning)
        if number is not None and number in recovered_numbers:
            continue
        result.append(warning)
    return result


async def _pro_ai_access(user_id: int) -> Tuple[bool, str]:
    if is_owner(user_id):
        return True, "owner"

    if not db.is_enabled():
        return False, "database"

    try:
        status = await db.get_plan_status(user_id, free_limit=FREE_IMPORT_LIMIT)
    except Exception:
        logging.exception("Could not read Pro status for AI")
        return False, "database"

    if not status.get("is_pro"):
        return False, "free"

    if not ai_parser.is_configured():
        return False, "not_configured"

    try:
        allowed, usage = await db.can_use_ai_import(
            user_id,
            monthly_limit=PRO_AI_IMPORT_LIMIT,
        )
    except Exception:
        logging.exception("Could not read AI usage")
        return False, "database"

    if not allowed:
        return False, "ai_limit"

    return True, "pro"


async def _run_ai_recovery(
    *,
    user_id: int,
    text: str,
    questions: List[dict],
    warnings: List[str],
    status_message,
) -> Tuple[List[dict], List[str], str, bool]:
    """
    Return (merged_questions, remaining_warnings, note, ai_called).
    AI failure never destroys deterministic-parser results.
    """
    allowed, reason = await _pro_ai_access(user_id)

    if not allowed:
        if reason == "free":
            return questions, warnings, "", False
        if reason == "ai_limit":
            return (
                questions,
                warnings,
                f"\n⚠️ PRO AI limiti: bu oy {PRO_AI_IMPORT_LIMIT} ta AI import ishlatildi.",
                False,
            )
        if reason == "not_configured":
            return questions, warnings, "\n⚠️ AI parser serverda sozlanmagan.", False
        return questions, warnings, "", False

    if not ai_parser.is_configured():
        return questions, warnings, "\n⚠️ GEMINI_API_KEY topilmadi.", False

    await status_message.edit_text(
        "🤖 Oddiy parser yetarli bo‘lmadi.\n"
        "PRO AI parser muammoli tuzilmani tekshirmoqda..."
    )

    try:
        result = await ai_parser.recover_questions(
            text=text,
            existing_questions=questions,
            parser_warnings=warnings,
        )
    except ai_parser.AIParserError as exc:
        code = str(exc)
        if code == "AI_QUOTA":
            note = "\n⚠️ Gemini bepul kvotasi hozircha tugagan yoki cheklangan."
        elif code == "AI_NOT_CONFIGURED":
            note = "\n⚠️ GEMINI_API_KEY sozlanmagan."
        elif code == "AI_MODEL_NOT_FOUND":
            note = "\n⚠️ Mos Gemini modeli topilmadi."
        else:
            note = "\n⚠️ AI parser hozir ishlamadi. Oddiy parser natijasi saqlandi."
        logging.warning("AI recovery failed: %s", code)
        return questions, warnings, note, False
    except Exception:
        logging.exception("Unexpected AI recovery failure")
        return (
            questions,
            warnings,
            "\n⚠️ AI parserda texnik xato bo‘ldi. Oddiy parser natijasi saqlandi.",
            False,
        )

    recovered = result.get("questions") or []
    merged = ai_parser.merge_questions(questions, recovered)
    remaining = _remove_recovered_warnings(warnings, recovered)

    for warning in result.get("warnings") or []:
        if warning not in remaining:
            remaining.append(warning)

    if recovered:
        note = f"\n🤖 AI parser: {len(recovered)} ta qo‘shimcha savol tiklandi."
    else:
        note = (
            "\n🤖 AI parser ishladi, lekin manbada aniq javobi "
            "tasdiqlangan qo‘shimcha savol topilmadi."
        )

    return merged, remaining, note, bool(result.get("ai_called"))


async def _run_scanned_pdf_recovery(*, user_id: int, raw: bytes, status_message):
    """Run Gemini vision/OCR only for PRO/owner when a PDF has no extractable text."""
    allowed, reason = await _pro_ai_access(user_id)
    if not allowed:
        if reason == "free":
            return [], [], "free", False
        if reason == "ai_limit":
            return [], [f"PRO AI limiti: bu oy {PRO_AI_IMPORT_LIMIT} ta AI import ishlatildi"], "ai_limit", False
        if reason == "not_configured":
            return [], ["AI parser serverda sozlanmagan"], "not_configured", False
        return [], ["AI parserga kirish tekshirilmadi"], reason, False

    await status_message.edit_text(
        "🖼 Matn qatlami topilmadi.\n"
        "👑 PRO AI skanerlangan PDF sahifalarini o‘qimoqda..."
    )
    try:
        result = await ai_parser.recover_scanned_pdf(raw)
    except ai_parser.AIParserError as exc:
        code = str(exc)
        if code.startswith("AI_SCAN_PAGE_LIMIT:"):
            parts = code.split(":")
            return [], [f"Skan PDF juda uzun: {parts[1]} sahifa. Hozirgi limit: {parts[2]} sahifa."], "page_limit", False
        if code == "AI_QUOTA":
            return [], ["Gemini kvotasi hozircha tugagan yoki cheklangan"], "quota", False
        if code in ("AI_NOT_CONFIGURED", "AI_MODEL_NOT_FOUND"):
            return [], ["AI skaner parseri hozir mavjud emas"], "unavailable", False
        logging.warning("Scanned PDF recovery failed: %s", code)
        return [], ["Skan PDF AI orqali o‘qilmadi"], "error", False
    except Exception:
        logging.exception("Unexpected scanned PDF AI failure")
        return [], ["Skan PDF AI orqali o‘qishda texnik xato bo‘ldi"], "error", False

    questions = result.get("questions") or []
    warnings = result.get("warnings") or []
    return questions, warnings, "ok", bool(result.get("ai_called"))


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    document = update.message.document
    user = update.effective_user
    user_id = user.id
    chat_id = update.effective_chat.id
    group_mode = is_group_chat(update.effective_chat)

    filename = document.file_name or "file"
    lower = filename.lower()

    if not (lower.endswith(".pdf") or lower.endswith(".docx")):
        await update.message.reply_text("❌ Hozircha faqat PDF va DOCX fayllar qabul qilinadi.")
        return

    if document.file_size and int(document.file_size) > MAX_UPLOAD_BYTES:
        await update.message.reply_text(
            f"❌ Fayl juda katta. Hozirgi limit: {MAX_UPLOAD_BYTES // (1024 * 1024)} MB."
        )
        return

    # In groups, an upload is part of the host's setup. Do not let a second user
    # overwrite another person's quiz session.
    if group_mode:
        ok, host_name = await claim_group_host_with_menu(chat_id, user, context.bot)
        if not ok:
            await update.message.reply_text(
                f"🔒 Bu guruh quizini hozir {host_name} boshqaryapti. "
                "Uning sessiyasi tugagach test yuklang."
            )
            return

    # Free plan: one NEW saved quiz import per calendar month.
    # Re-uploading the same filename is treated as an update and does not consume
    # another slot. OWNER_TELEGRAM_ID bypasses this during private testing.
    if db.is_enabled() and not is_owner(user_id):
        try:
            await db.ensure_user(user_id, user.username, user.full_name)
            allowed, reason = await db.can_import_new_quiz(
                user_id,
                filename,
                free_limit=FREE_IMPORT_LIMIT,
            )
            if not allowed:
                await update.message.reply_text(
                    "🔒 Bepul tarifda oyiga 1 ta yangi test import qilish mumkin.\n\n"
                    "Oldin saqlangan testlaringizni cheksiz ishlashingiz mumkin. "
                    "Yangi bepul import keyingi oyda ochiladi."
                )
                return
        except Exception:
            logging.exception("Could not verify monthly import limit")

    status = await update.message.reply_text("📥 Fayl o‘qilmoqda...")

    try:
        tg_file = await context.bot.get_file(document.file_id)
        data = await tg_file.download_as_bytearray()
        raw = bytes(data)

        parser_method = "standart parser"
        ai_note = ""
        questions = []
        warnings = []
        ai_used_this_upload = False
        ai_recovered_count = 0

        if lower.endswith(".pdf"):
            # Deterministic support for teacher PDFs where answers are marked
            # with real PDF Highlight annotations inside answer-table cells.
            table_questions, table_warnings = parser_engine.parse_highlighted_pdf_tables(raw)
            table_total = len(table_questions) + len(table_warnings)
            table_ratio = (len(table_questions) / table_total) if table_total else 0.0
            if len(table_questions) >= 5 and table_ratio >= 0.80:
                questions, warnings = table_questions, table_warnings
                text = parser_engine.extract_pdf_text(raw)
                parser_method = "PDF highlight-jadval parseri"
            else:
                text = parser_engine.extract_pdf_text(raw)
        else:
            text = parser_engine.extract_docx_text(raw)

        if not questions and not text.strip():
            if lower.endswith(".pdf"):
                questions, warnings, scan_status, scan_ai_used = await _run_scanned_pdf_recovery(
                    user_id=user_id, raw=raw, status_message=status
                )
                ai_used_this_upload = ai_used_this_upload or scan_ai_used
                if scan_ai_used:
                    ai_recovered_count += len(questions)
                if scan_status == "free":
                    await status.edit_text(
                        "❌ Bu PDF skaner/rasm ko‘rinishida. Matn qatlami topilmadi.\n\n"
                        "👑 PRO da skanerlangan PDF AI orqali o‘qiladi."
                    )
                    return
                if not questions:
                    detail = "\n".join(f"• {w}" for w in warnings[:8])
                    await status.edit_text(
                        "❌ Skan PDFdan quizga tayyor savol olinmadi."
                        + (f"\n\n{detail}" if detail else "")
                    )
                    return
                parser_method = "PRO AI skan parseri"
                ai_note = f"\n🖼 PRO AI scan: {len(questions)} ta savol tayyorlandi."
            else:
                await status.edit_text("❌ DOCX fayldan o‘qiladigan matn topilmadi.")
                return
        elif not questions:
            await status.edit_text("🔎 Test savollari aniqlanmoqda...")
            questions, warnings = parser_engine.parse_questions(text)

            # PRO AI fallback runs automatically only when deterministic parsing
            # leaves unresolved blocks or finds no ready questions.
            if warnings or not questions:
                before_ai = len(questions)
                questions, warnings, ai_note, text_ai_used = await _run_ai_recovery(
                    user_id=user_id,
                    text=text,
                    questions=questions,
                    warnings=warnings,
                    status_message=status,
                )
                ai_used_this_upload = ai_used_this_upload or text_ai_used
                ai_recovered_count += max(0, len(questions) - before_ai)
                if len(questions) > before_ai:
                    parser_method = "standart + PRO AI parser"

            # Some scanned PDFs contain a tiny/garbled OCR text layer. If both
            # deterministic and text-AI recovery still produce nothing, let PRO
            # users fall back to the visual PDF reader as a final attempt.
            if not questions and lower.endswith(".pdf"):
                scan_questions, scan_warnings, scan_status, scan_ai_used = await _run_scanned_pdf_recovery(
                    user_id=user_id, raw=raw, status_message=status
                )
                ai_used_this_upload = ai_used_this_upload or scan_ai_used
                if scan_ai_used:
                    ai_recovered_count += len(scan_questions)
                if scan_questions:
                    questions, warnings = scan_questions, scan_warnings
                    parser_method = "PRO AI skan parseri"
                    ai_note += f"\n🖼 PRO AI scan: {len(questions)} ta savol tayyorlandi."

        total_blocks = len(questions) + len(warnings)

        if not questions:
            if ai_used_this_upload and db.is_enabled() and not is_owner(user_id):
                try:
                    await db.record_ai_import(user_id, ai_recovered_count)
                except Exception:
                    logging.exception("Could not record AI usage")
            ai_access, ai_reason = await _pro_ai_access(user_id)
            pro_hint = ""
            if ai_reason == "free":
                pro_hint = (
                    "\n\n👑 PRO foydalanuvchilarda noodatiy matnli PDF/DOCX "
                    "formatlar avtomatik AI parser orqali tekshiriladi."
                )

            await status.edit_text(
                "❌ Quizga tayyor savol topilmadi.\n\n"
                "Qo‘llab-quvvatlanadigan asosiy format:\n\n"
                "1. Savol\nA) ...\nB) ...\nC) ...\nD) ...\nJavob: B\n\n"
                "Bot to‘g‘ri javobni taxmin qilmaydi."
                f"{ai_note}{pro_hint}"
            )
            return

        if ai_used_this_upload and db.is_enabled() and not is_owner(user_id):
            try:
                await db.record_ai_import(user_id, ai_recovered_count)
            except Exception:
                logging.exception("Could not record AI usage")

        prefs = await effective_preferences(user_id)
        session_data = {
            "chat_id": chat_id,
            "filename": filename,
            "questions": questions,
            "warnings": warnings,
            "group_size": None,
            "groups": [],
            "active": None,
            "results": {},
            "controller_id": user_id,
            "controller_name": user.full_name,
            "owner_username": user.username,
            "owner_full_name": user.full_name,
            "host_last_activity": time.time(),
            "last_leaderboard_text": None,
            "shuffle_questions": bool(prefs.get("shuffle_questions")),
            "shuffle_options": bool(prefs.get("shuffle_options")),
            "quiz_mode": None,
            "parser_total_blocks": total_blocks,
            "parser_method": parser_method,
            "default_group_size": int(prefs.get("default_group_size", 50)),
            "default_timer": int(prefs.get("default_timer", 30)),
        }

        # Save first so the RAM session knows the persistent quiz id.
        saved_note = ""
        if db.is_enabled():
            try:
                save_result = await db.save_quiz(
                    owner_id=user_id,
                    username=user.username,
                    full_name=user.full_name,
                    filename=filename,
                    questions=questions,
                )
                if save_result:
                    session_data["saved_quiz_id"] = int(save_result["quiz_id"])
                    if save_result.get("created_new"):
                        saved_note = "\n💾 Yangi test bazaga saqlandi."
                    else:
                        saved_note = "\n💾 Saqlangan test yangilandi."
            except Exception:
                logging.exception("Could not save quiz to database")
                saved_note = "\n⚠️ Test ishlaydi, lekin bazaga saqlashda xato bo‘ldi."

        if group_mode:
            GROUP_DATA[chat_id] = session_data
        else:
            USER_DATA[user_id] = session_data

        warning_text = ""
        if warnings:
            preview = "\n".join(f"• {w}" for w in warnings[:8])
            more = f"\n• ... yana {len(warnings) - 8} ta" if len(warnings) > 8 else ""
            warning_text = (
                f"\n⚠️ Muammoli savollar: {len(warnings)}\n"
                f"{preview}{more}"
            )

        mode_note = ""
        if group_mode:
            mode_note = (
                f"\n👤 Boshqaruvchi: {user.full_name}\n"
                "Faqat boshqaruvchi setup/START/skip/stop tugmalarini ishlata oladi."
            )

        await status.edit_text(
            f"✅ Fayl tahlil qilindi.\n\n"
            f"📄 {filename}\n"
            f"🧩 Savol bloklari topildi: {total_blocks}\n"
            f"✅ Quizga tayyor: {len(questions)}\n"
            f"🔧 Usul: {parser_method}"
            f"{warning_text}"
            f"{mode_note}"
            f"{ai_note}"
            f"{saved_note}\n\n"
            "Har bir bo‘limda nechta savol bo‘lsin?",
            reply_markup=group_size_keyboard(group_mode=group_mode, default_size=session_data.get("default_group_size")),
        )

    except Exception as e:
        logging.exception("Document processing failed")
        await status.edit_text(f"❌ Faylni qayta ishlashda xato yuz berdi:\n{type(e).__name__}")


def build_groups(questions: List[dict], size: int) -> List[List[dict]]:
    return [questions[i:i + size] for i in range(0, len(questions), size)]


async def choose_size_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    session = USER_DATA.get(user_id)
    if not session:
        await query.message.reply_text("❌ Sessiya topilmadi. Faylni qayta yuboring.")
        return

    size = int(query.data.split(":")[1])
    session.pop("custom_questions", None)
    session.pop("custom_range_label", None)
    session["group_size"] = size
    session.setdefault("shuffle_questions", False)
    session.setdefault("shuffle_options", False)
    await send_order_settings(query.message, session, group_mode=False)


async def custom_range_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    user_id = update.effective_user.id
    group_mode = query.data == "gcustomrange"

    if not await user_has_pro(user_id):
        await query.answer("🎯 Maxsus savol oralig‘i PRO funksiyasi.", show_alert=True)
        return

    if group_mode:
        session = GROUP_DATA.get(update.effective_chat.id)
        if not session or not group_controller_ok(update, session):
            await query.answer("Faqat guruh boshqaruvchisi tanlaydi.", show_alert=True)
            return
    else:
        session = USER_DATA.get(user_id)
        if not session:
            await query.answer("Sessiya topilmadi.", show_alert=True)
            return

    await query.answer()
    context.user_data["custom_range_request"] = {
        "mode": "group" if group_mode else "private",
        "chat_id": update.effective_chat.id,
    }
    await query.message.reply_text(
        f"🎯 Savollar oralig‘ini yozing. Jami: {len(session.get('questions', []))} ta.\n\n"
        "Masalan: 121-170\n"
        "Bekor qilish uchun /cancel yuboring."
    )


async def apply_custom_range_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> bool:
    request = context.user_data.get("custom_range_request")
    if not request:
        return False

    user_id = update.effective_user.id
    mode = request.get("mode")
    if mode == "group":
        session = GROUP_DATA.get(update.effective_chat.id)
        if not session or session.get("controller_id") != user_id:
            context.user_data.pop("custom_range_request", None)
            await update.message.reply_text("❌ Guruh sessiyasi topilmadi yoki boshqaruvchi o‘zgargan.")
            return True
    else:
        session = USER_DATA.get(user_id)
        if not session:
            context.user_data.pop("custom_range_request", None)
            await update.message.reply_text("❌ Sessiya topilmadi.")
            return True

    bounds = pro_features.parse_custom_range(
        update.message.text or "",
        len(session.get("questions", [])),
    )
    if not bounds:
        await update.message.reply_text(
            "❌ Oraliq noto‘g‘ri. Masalan: 121-170. "
            f"Jami savollar: {len(session.get('questions', []))}."
        )
        return True

    start, end = bounds
    session["custom_questions"] = list(session["questions"][start - 1:end])
    session["custom_range_label"] = pro_features.range_label(start, end)
    session["group_size"] = len(session["custom_questions"])
    context.user_data.pop("custom_range_request", None)

    await update.message.reply_text(
        f"✅ Maxsus oraliq tanlandi: {start}–{end} ({end - start + 1} ta savol)."
    )
    await send_order_settings(update.message, session, group_mode=(mode == "group"))
    return True


async def custom_random_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    user_id = update.effective_user.id
    group_mode = query.data == "gcustomrandom"

    if not await user_has_pro(user_id):
        await query.answer("🎲 Tasodifiy N ta savol PRO funksiyasi.", show_alert=True)
        return

    if group_mode:
        session = GROUP_DATA.get(update.effective_chat.id)
        if not session or not group_controller_ok(update, session):
            await query.answer("Faqat guruh boshqaruvchisi tanlaydi.", show_alert=True)
            return
    else:
        session = USER_DATA.get(user_id)
        if not session:
            await query.answer("Sessiya topilmadi.", show_alert=True)
            return

    await query.answer()
    context.user_data["custom_random_request"] = {
        "mode": "group" if group_mode else "private",
        "chat_id": update.effective_chat.id,
    }
    await query.message.reply_text(
        f"🎲 Nechta tasodifiy savol kerak? Jami: {len(session.get('questions', []))} ta.\n\n"
        "Masalan: 75\n"
        "Bekor qilish uchun /cancel yuboring."
    )


async def apply_custom_random_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> bool:
    request = context.user_data.get("custom_random_request")
    if not request:
        return False

    user_id = update.effective_user.id
    mode = request.get("mode")
    if mode == "group":
        session = GROUP_DATA.get(update.effective_chat.id)
        if not session or session.get("controller_id") != user_id:
            context.user_data.pop("custom_random_request", None)
            await update.message.reply_text("❌ Guruh sessiyasi topilmadi yoki boshqaruvchi o‘zgargan.")
            return True
    else:
        session = USER_DATA.get(user_id)
        if not session:
            context.user_data.pop("custom_random_request", None)
            await update.message.reply_text("❌ Sessiya topilmadi.")
            return True

    total = len(session.get("questions", []))
    count = pro_features.parse_random_count(update.message.text or "", total)
    if count is None:
        await update.message.reply_text(
            f"❌ Son noto‘g‘ri. 1 dan {total} gacha butun son yozing. Masalan: 75."
        )
        return True

    selected = random.sample(list(session["questions"]), count)
    session["custom_questions"] = selected
    session["custom_range_label"] = f"Tasodifiy {count} ta"
    session["group_size"] = count
    context.user_data.pop("custom_random_request", None)

    await update.message.reply_text(f"✅ {count} ta tasodifiy savol tanlandi.")
    await send_order_settings(update.message, session, group_mode=(mode == "group"))
    return True


async def private_toggle_questions_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    session = USER_DATA.get(update.effective_user.id)
    if not session:
        await query.answer("Sessiya topilmadi", show_alert=True)
        return
    session["shuffle_questions"] = not bool(session.get("shuffle_questions"))
    await query.answer()
    await query.edit_message_reply_markup(
        reply_markup=order_settings_keyboard(session, group_mode=False)
    )


async def private_toggle_options_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    session = USER_DATA.get(update.effective_user.id)
    if not session:
        await query.answer("Sessiya topilmadi", show_alert=True)
        return
    session["shuffle_options"] = not bool(session.get("shuffle_options"))
    await query.answer()
    await query.edit_message_reply_markup(
        reply_markup=order_settings_keyboard(session, group_mode=False)
    )


async def private_toggle_mode_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    session = USER_DATA.get(update.effective_user.id)
    if not session:
        return
    session["quiz_mode"] = "exam" if session.get("quiz_mode", "practice") == "practice" else "practice"
    await query.edit_message_reply_markup(reply_markup=order_settings_keyboard(session, False))


async def private_order_done_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    session = USER_DATA.get(user_id)
    if not session or not session.get("group_size"):
        await query.message.reply_text("❌ Avval guruh hajmini tanlang.")
        return
    ordered = apply_question_order(session)
    session["ordered_questions"] = ordered
    session["groups"] = build_groups(ordered, session["group_size"])
    await show_groups(query.message, user_id)


async def show_groups(message, user_id: int):
    session = USER_DATA.get(user_id)
    if not session:
        await message.reply_text("❌ Sessiya topilmadi.")
        return

    rows = []
    for idx, group in enumerate(session["groups"]):
        start_no = idx * session["group_size"] + 1
        end_no = start_no + len(group) - 1

        result = session["results"].get(idx)
        suffix = ""
        if result:
            suffix = f" ✅ {result['percent']}%"

        if session.get("custom_range_label") and len(session.get("groups", [])) == 1:
            label = f"🎯 {session['custom_range_label']} · {len(group)} ta{suffix}"
        elif session.get("shuffle_questions"):
            label = f"📘 {idx + 1}-guruh · {len(group)} ta 🔀{suffix}"
        else:
            label = f"📘 {idx + 1}-guruh · {start_no}-{end_no}{suffix}"
        rows.append(
            [
                InlineKeyboardButton(
                    label,
                    callback_data=f"group:{idx}",
                )
            ]
        )

    rows.append([InlineKeyboardButton("📄 Boshqa fayl yuborish", callback_data="help_upload")])
    rows.append([InlineKeyboardButton("🏠 Bosh menyu", callback_data="menu_home")])

    await message.reply_text(
        f"✅ {len(session['groups'])} ta guruh tayyor.\n\n"
        "Guruhni tanlang. Quiz faqat ▶️ Start bosilgandan keyin boshlanadi:",
        reply_markup=InlineKeyboardMarkup(rows),
    )


async def group_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = update.effective_user.id
    session = USER_DATA.get(user_id)
    if not session:
        await query.message.reply_text("❌ Sessiya topilmadi.")
        return

    group_index = int(query.data.split(":")[1])
    if group_index < 0 or group_index >= len(session["groups"]):
        return

    group = session["groups"][group_index]
    start_no = group_index * session["group_size"] + 1
    end_no = start_no + len(group) - 1

    previous = session["results"].get(group_index)
    previous_text = ""
    if previous:
        previous_text = (
            f"\nOldingi natija: {previous['correct']}/{previous['total']} "
            f"({previous['percent']}%)"
        )

    keyboard = timer_keyboard(group_index, group_mode=False, default_timer=session.get("default_timer"))

    range_text = (
        f"🎯 Savollar: {session['custom_range_label']}"
        if session.get("custom_range_label") and len(session.get("groups", [])) == 1
        else ("🔀 Aralashtirilgan savollar" if session.get("shuffle_questions") else f"Savollar: {start_no}-{end_no}")
    )
    await query.message.reply_text(
        f"📘 {group_index + 1}-guruh\n"
        f"{range_text}\n"
        f"Jami: {len(group)}"
        f"{previous_text}\n\n"
        "⏱ Har bir savol uchun vaqtni tanlang:",
        reply_markup=keyboard,
    )


async def timer_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = update.effective_user.id
    session = USER_DATA.get(user_id)
    if not session:
        await query.message.reply_text("❌ Sessiya topilmadi.")
        return

    _, group_text, seconds_text = query.data.split(":")
    group_index = int(group_text)
    seconds = int(seconds_text)

    if group_index < 0 or group_index >= len(session["groups"]):
        return
    if seconds not in TIMER_CHOICES:
        return

    session["selected_timer"] = seconds
    session["quiz_mode"] = None
    session["pending_start"] = {
        "group_index": group_index,
        "timer_seconds": seconds,
        "mode": None,
    }

    await query.message.reply_text(
        "🎮 TEST REJIMINI TANLANG\n\n"
        "📖 Mashq rejimi\n"
        "• Har savoldan keyin Telegram to‘g‘ri/noto‘g‘ri javobni ko‘rsatadi.\n"
        "• Xatolarni keyin qayta mashq qilish mumkin.\n\n"
        "📝 Imtihon rejimi\n"
        "• Savol paytida to‘g‘ri javob ko‘rsatilmaydi.\n"
        "• Natija test oxirida chiqadi.\n\n"
        "Davom etish uchun bittasini tanlang:",
        reply_markup=InlineKeyboardMarkup(
            [[
                InlineKeyboardButton(
                    "📖 Mashq rejimi",
                    callback_data=f"pmode:{group_index}:{seconds}:practice",
                ),
                InlineKeyboardButton(
                    "📝 Imtihon rejimi",
                    callback_data=f"pmode:{group_index}:{seconds}:exam",
                ),
            ],[
                InlineKeyboardButton(
                    "⏱ Vaqtni o‘zgartirish",
                    callback_data=f"group:{group_index}",
                )
            ]]
        ),
    )


async def private_mode_choice_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = update.effective_user.id
    session = USER_DATA.get(user_id)
    if not session:
        await query.message.reply_text("❌ Sessiya topilmadi.")
        return

    _, group_text, seconds_text, mode = query.data.split(":")
    group_index = int(group_text)
    seconds = int(seconds_text)

    if mode not in ("practice", "exam"):
        return
    if seconds not in TIMER_CHOICES:
        return
    if group_index < 0 or group_index >= len(session.get("groups", [])):
        return

    session["quiz_mode"] = mode
    session["pending_start"] = {
        "group_index": group_index,
        "timer_seconds": seconds,
        "mode": mode,
    }

    group = session["groups"][group_index]
    max_seconds = len(group) * seconds
    max_time_text = (
        f"{max_seconds / 60:.1f} daqiqagacha"
        if max_seconds >= 60
        else f"{max_seconds} soniyagacha"
    )
    mode_text = "📖 Mashq rejimi" if mode == "practice" else "📝 Imtihon rejimi"

    await query.message.reply_text(
        f"✅ Test tayyor.\n\n"
        f"📘 Guruh: {group_index + 1}\n"
        f"❓ Savollar: {len(group)}\n"
        f"⏱ Har bir savol: {format_duration(seconds)}\n"
        f"🎮 Rejim: {mode_text}\n"
        f"⌛ Maksimal vaqt: {max_time_text}\n\n"
        "Test hali boshlanmadi. Boshlash uchun ▶️ START ni bosing.",
        reply_markup=InlineKeyboardMarkup(
            [
                [InlineKeyboardButton(
                    "▶️ START",
                    callback_data=f"startgroup:{group_index}:{seconds}",
                )],
                [InlineKeyboardButton(
                    "🎮 Rejimni o‘zgartirish",
                    callback_data=f"timer:{group_index}:{seconds}",
                )],
                [InlineKeyboardButton(
                    "⏱ Vaqtni o‘zgartirish",
                    callback_data=f"group:{group_index}",
                )],
                [InlineKeyboardButton("📚 Guruhlar", callback_data="groups")],
            ]
        ),
    )


async def groups_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    await show_groups(query.message, update.effective_user.id)


async def start_group_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = update.effective_user.id
    chat_id = update.effective_chat.id
    session = USER_DATA.get(user_id)

    if not session:
        await query.message.reply_text("❌ Sessiya topilmadi.")
        return
    if session.get("active"):
        await query.message.reply_text("⚠️ Quiz allaqachon ishlayapti.")
        return

    parts = query.data.split(":")
    if len(parts) != 3:
        await query.message.reply_text("❌ Vaqt tanlanmagan. Guruhni qayta tanlang.")
        return

    group_index = int(parts[1])
    timer_seconds = int(parts[2])

    if group_index < 0 or group_index >= len(session["groups"]):
        return
    if timer_seconds not in TIMER_CHOICES:
        return

    pending = session.get("pending_start") or {}
    if (
        pending.get("group_index") != group_index
        or pending.get("timer_seconds") != timer_seconds
        or pending.get("mode") not in ("practice", "exam")
    ):
        await query.message.reply_text("🎮 Avval Mashq yoki Imtihon rejimini tanlang.")
        return
    session["quiz_mode"] = pending["mode"]

    # A run_id prevents an old timeout task from affecting a restarted quiz.
    session["run_counter"] = session.get("run_counter", 0) + 1
    run_id = session["run_counter"]

    session["active"] = {
        "run_id": run_id,
        "group_index": group_index,
        "questions": session["groups"][group_index],
        "current": 0,
        "correct": 0,
        "wrong": [],
        "unanswered": [],
        "answered_polls": set(),
        "timer_seconds": timer_seconds,
        "current_streak": 0,
        "best_streak": 0,
        "review_mode": False,
        "empty_streak": 0,
        "paused": False,
        "current_poll_id": None,
        "current_poll_message_id": None,
        "quiz_mode": session.get("quiz_mode", "practice"),
        "question_outcomes": {},
        "special_mode": session.get("special_mode"),
        "pro_features": await user_has_pro(user_id),
        "bookmarked_positions": set(),
    }
    if session["active"]["pro_features"] and session.get("saved_quiz_id") and db.is_enabled():
        try:
            session["active"]["bookmarked_positions"] = set(
                await db.list_bookmark_positions(user_id, int(session["saved_quiz_id"]))
            )
        except Exception:
            logging.exception("Bookmarklar yuklanmadi")

    try:
        await context.bot.send_dice(chat_id=chat_id, emoji="🎯")
    except Exception:
        pass

    await query.message.reply_text(
        f"🚀 {group_index + 1}-guruh boshlandi!\n"
        f"Jami {len(session['active']['questions'])} ta savol.\n"
        f"⏱ Har bir savol uchun: {format_duration(timer_seconds)}\n"
        f"🔀 Savollar: {'yoqilgan' if session.get('shuffle_questions') else 'o‘chirilgan'} · "
        f"Variantlar: {'yoqilgan' if session.get('shuffle_options') else 'o‘chirilgan'}\n"
        f"🎮 Rejim: {'Imtihon' if session.get('quiz_mode') == 'exam' else 'Mashq'}"
    )
    await send_next_question(chat_id, user_id, context)


def telegram_safe_question(text: str) -> str:
    # Telegram poll question limit is finite; keep it safely short.
    text = re.sub(r"\s+", " ", text).strip()
    return text[:290] if len(text) > 290 else text


def telegram_safe_option(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    return text[:95] if len(text) > 95 else text


def private_bookmark_markup(session: dict, active: dict, item: dict) -> Optional[InlineKeyboardMarkup]:
    """Build the persistent bookmark keyboard for one private Pro poll."""
    if not active.get("pro_features") or not session.get("saved_quiz_id"):
        return None

    try:
        position = int(item.get("position") or 0)
        quiz_id = int(session["saved_quiz_id"])
        run_id = int(active["run_id"])
    except Exception:
        return None

    if position <= 0:
        return None

    marked = position in active.get("bookmarked_positions", set())
    return InlineKeyboardMarkup([[
        InlineKeyboardButton(
            "🌟 Belgilangan" if marked else "⭐ Keyin ko‘rish",
            callback_data=f"pbookmarktoggle:{quiz_id}:{run_id}:{position}",
        )
    ]])


def bookmark_preview_text(quiz: dict, positions: List[int], max_chars: int = 3300) -> str:
    """Show the actual bookmarked questions while staying under Telegram limits."""
    by_pos = {
        int(q.get("position") or 0): q
        for q in quiz.get("questions", [])
        if int(q.get("position") or 0) > 0
    }

    lines = [f"⭐ Belgilangan savollar: {len(positions)} ta", ""]
    shown = 0

    for pos in positions:
        q = by_pos.get(int(pos))
        if not q:
            continue

        question = re.sub(r"\s+", " ", str(q.get("question") or "")).strip()
        if len(question) > 180:
            question = question[:177].rstrip() + "..."

        source_no = q.get("number") or pos
        line = f"⭐ {source_no}. {question}"

        projected = "\n\n".join(lines + [line])
        if len(projected) > max_chars:
            break

        lines.append(line)
        shown += 1

    remaining = max(0, len(positions) - shown)
    if remaining:
        lines += ["", f"… yana {remaining} ta belgilangan savol."]

    return "\n\n".join(lines)


async def send_next_question(chat_id: int, user_id: int, context: ContextTypes.DEFAULT_TYPE):
    session = USER_DATA.get(user_id)
    if not session or not session.get("active"):
        return

    active = session["active"]

    if active.get("paused"):
        return

    idx = active["current"]
    questions = active["questions"]

    if idx >= len(questions):
        await finish_group(chat_id, user_id, context)
        return

    item = questions[idx]
    displayed_options, displayed_correct_index = prepare_poll_options(
        item,
        bool(session.get("shuffle_options")),
    )
    options = [telegram_safe_option(x) for x in displayed_options]
    timer_seconds = active["timer_seconds"]

    try:
        poll_kwargs = dict(
            chat_id=chat_id,
            question=telegram_safe_question(f"[{idx + 1}/{len(questions)}] {item['question']}"),
            options=options,
            is_anonymous=False,
            allows_multiple_answers=False,
            open_period=timer_seconds,
        )
        if active.get("quiz_mode", "practice") == "exam":
            poll_kwargs["type"] = "regular"
        else:
            poll_kwargs["type"] = "quiz"
            poll_kwargs["correct_option_id"] = displayed_correct_index

        bookmark_markup = private_bookmark_markup(session, active, item)
        if bookmark_markup is not None:
            poll_kwargs["reply_markup"] = bookmark_markup
        msg = await context.bot.send_poll(**poll_kwargs)
    except Exception:
        logging.exception("Could not send poll")
        await context.bot.send_message(
            chat_id=chat_id,
            text="❌ Bu savol Telegram formatiga sig‘madi. Keyingi savolga o‘tyapman.",
        )
        active["wrong"].append(idx)
        active["current"] += 1
        await send_next_question(chat_id, user_id, context)
        return

    active["current_poll_id"] = msg.poll.id
    active["current_poll_message_id"] = msg.message_id

    POLL_MAP[msg.poll.id] = {
        "mode": "private",
        "user_id": user_id,
        "chat_id": chat_id,
        "message_id": msg.message_id,
        "group_index": active["group_index"],
        "question_index": idx,
        "run_id": active["run_id"],
        "poll_correct_index": displayed_correct_index,
        "bookmark_position": int(item.get("position") or 0),
        "quiz_id": int(session["saved_quiz_id"]) if session.get("saved_quiz_id") else None,
        "handled": False,
    }

    # PollAnswerHandler is not called when the user gives no answer.
    # This task moves the quiz forward once the poll's timer expires.
    asyncio.create_task(
        question_timeout(
            poll_id=msg.poll.id,
            user_id=user_id,
            chat_id=chat_id,
            group_index=active["group_index"],
            question_index=idx,
            run_id=active["run_id"],
            timer_seconds=timer_seconds,
            context=context,
        )
    )


async def bookmark_toggle_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    user_id = update.effective_user.id

    if not await user_has_pro(user_id):
        await query.answer("Bu PRO funksiyasi.", show_alert=True)
        return
    if not db.is_enabled():
        await query.answer("Database ulanmagan.", show_alert=True)
        return

    _, quiz_text, run_text, pos_text = query.data.split(":")
    quiz_id = int(quiz_text)
    position = int(pos_text)

    # Do NOT require the old quiz to still be active in RAM.
    # Ownership and validity are checked by the persistent database.
    try:
        marked = await db.toggle_bookmark(user_id, quiz_id, position)
    except Exception:
        logging.exception("Bookmark toggle failed")
        await query.answer("Belgilashni saqlab bo‘lmadi.", show_alert=True)
        return

    # Keep current active-session state in sync only when it is the same saved quiz.
    session = USER_DATA.get(user_id)
    active = session.get("active") if session else None
    if (
        active
        and session.get("saved_quiz_id")
        and int(session["saved_quiz_id"]) == quiz_id
    ):
        marks = active.setdefault("bookmarked_positions", set())
        if marked:
            marks.add(position)
        else:
            marks.discard(position)

    await query.answer(
        "⭐ Keyin ko‘rish uchun saqlandi"
        if marked else
        "⭐ Belgilanganlardan olib tashlandi"
    )

    try:
        await query.edit_message_reply_markup(
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton(
                    "🌟 Belgilangan" if marked else "⭐ Keyin ko‘rish",
                    callback_data=f"pbookmarktoggle:{quiz_id}:{run_text}:{position}",
                )
            ]])
        )
    except Exception:
        # A closed/old poll may reject markup edits on some Telegram clients.
        # The bookmark is already safely persisted in DB.
        pass


async def question_timeout(
    poll_id: str,
    user_id: int,
    chat_id: int,
    group_index: int,
    question_index: int,
    run_id: int,
    timer_seconds: int,
    context: ContextTypes.DEFAULT_TYPE,
):
    # Small buffer lets Telegram close the poll first.
    await asyncio.sleep(timer_seconds + 0.8)

    meta = POLL_MAP.get(poll_id)
    if not meta or meta.get("handled"):
        return

    session = USER_DATA.get(user_id)
    if not session or not session.get("active"):
        POLL_MAP.pop(poll_id, None)
        return

    active = session["active"]

    # Ignore timeout tasks from an older/restarted quiz.
    if (
        active.get("run_id") != run_id
        or active.get("group_index") != group_index
        or active.get("current") != question_index
    ):
        POLL_MAP.pop(poll_id, None)
        return

    # Mark this poll before awaiting anything else, preventing a late answer
    # from advancing the quiz twice.
    meta["handled"] = True
    active["unanswered"].append(question_index)
    try:
        pos = int(active["questions"][question_index].get("position") or 0)
        if pos:
            active.setdefault("question_outcomes", {})[pos] = "unanswered"
    except Exception:
        pass
    active["current_streak"] = 0
    active["empty_streak"] = active.get("empty_streak", 0) + 1
    active["current"] += 1
    active["current_poll_id"] = None
    active["current_poll_message_id"] = None
    POLL_MAP.pop(poll_id, None)

    if active["empty_streak"] >= 3:
        active["paused"] = True
        await context.bot.send_message(
            chat_id=chat_id,
            text=(
                "⏸ Quiz avtomatik pauzaga qo‘yildi.\n\n"
                "Oxirgi 3 ta savol javobsiz qoldi. "
                "Davom ettirish uchun /resume yuboring yoki butunlay tugatish uchun /stop yuboring."
            ),
        )
        return

    await context.bot.send_message(
        chat_id=chat_id,
        text=f"⏱ Vaqt tugadi. {question_index + 1}-savol javobsiz qoldi.",
    )

    await asyncio.sleep(QUESTION_TRANSITION_DELAY)
    await send_next_question(chat_id, user_id, context)


async def poll_answer_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answer = update.poll_answer
    meta = POLL_MAP.get(answer.poll_id)
    if not meta:
        return

    if meta.get("mode") == "group":
        await group_poll_answer_handler(update, context, meta)
        return

    # Only count the user who started this private quiz.
    if answer.user.id != meta["user_id"]:
        return

    # If timeout already handled this poll, ignore any late update.
    if meta.get("handled"):
        return

    user_id = meta["user_id"]
    session = USER_DATA.get(user_id)
    if not session or not session.get("active"):
        return

    active = session["active"]

    # Ignore stale answers from an older/restarted quiz.
    if (
        active.get("run_id") != meta.get("run_id")
        or active["group_index"] != meta["group_index"]
    ):
        POLL_MAP.pop(answer.poll_id, None)
        return

    q_idx = meta["question_index"]
    if q_idx != active["current"]:
        return

    if answer.poll_id in active["answered_polls"]:
        return

    # Mark handled BEFORE awaiting the next question. This prevents the timer
    # task from advancing the quiz at the same time.
    meta["handled"] = True
    active["answered_polls"].add(answer.poll_id)

    item = active["questions"][q_idx]
    selected = answer.option_ids[0] if answer.option_ids else None
    correct_index = int(meta.get("poll_correct_index", item["correct_index"]))

    position = int(item.get("position") or 0)
    if selected == correct_index:
        if position:
            active.setdefault("question_outcomes", {})[position] = "correct"
        active["correct"] += 1
        active["current_streak"] = active.get("current_streak", 0) + 1
        active["best_streak"] = max(
            active.get("best_streak", 0),
            active["current_streak"],
        )

        if active["current_streak"] in {5, 10, 20, 30, 50}:
            await context.bot.send_message(
                chat_id=meta["chat_id"],
                text=f"🔥 {active['current_streak']} ta ketma-ket to‘g‘ri!",
            )
    else:
        if position:
            active.setdefault("question_outcomes", {})[position] = "wrong"
        active["wrong"].append(q_idx)
        active["current_streak"] = 0

    # Any submitted answer means the user is active again.
    active["empty_streak"] = 0
    active["current"] += 1
    message_id = meta.get("message_id")
    if message_id:
        try:
            persistent_markup = private_bookmark_markup(session, active, item)
            await context.bot.stop_poll(
                meta["chat_id"],
                message_id,
                reply_markup=persistent_markup,
            )
        except Exception:
            logging.exception("Answered poll could not be stopped with bookmark keyboard")
            try:
                await context.bot.stop_poll(meta["chat_id"], message_id)
            except Exception:
                pass
    active["current_poll_id"] = None
    active["current_poll_message_id"] = None
    POLL_MAP.pop(answer.poll_id, None)

    await asyncio.sleep(QUESTION_TRANSITION_DELAY)
    await send_next_question(meta["chat_id"], user_id, context)


async def finish_group(chat_id: int, user_id: int, context: ContextTypes.DEFAULT_TYPE, completed_count: Optional[int] = None, stopped_reason: Optional[str] = None):
    session = USER_DATA[user_id]
    active = session["active"]

    total = completed_count if completed_count is not None else len(active["questions"])
    correct = active["correct"]
    wrong_answered = len(active["wrong"])
    unanswered = len(active["unanswered"])
    percent = round((correct / total) * 100) if total else 0
    group_index = active["group_index"]
    timer_seconds = active["timer_seconds"]

    best_streak = active.get("best_streak", 0)
    problem_indices = list(dict.fromkeys(active["wrong"] + active["unanswered"]))
    problem_questions = [
        active["questions"][i]
        for i in problem_indices
        if 0 <= i < len(active["questions"])
    ]

    # If a review was manually stopped early, questions not reached yet are
    # still unresolved and must remain in the next review round.
    if active.get("review_mode") and completed_count is not None:
        unresolved_tail = active["questions"][completed_count:]
        seen_ids = {id(q) for q in problem_questions}
        for q in unresolved_tail:
            if id(q) not in seen_ids:
                problem_questions.append(q)
                seen_ids.add(id(q))

    if not active.get("review_mode"):
        session["results"][group_index] = {
            "correct": correct,
            "wrong": wrong_answered,
            "unanswered": unanswered,
            "total": total,
            "percent": percent,
            "timer_seconds": timer_seconds,
            "best_streak": best_streak,
            "problem_questions": problem_questions,
        }
    else:
        # Review mode is progressive: only mistakes from THIS review survive.
        # Corrected questions disappear from the next review round.
        previous_result = session["results"].setdefault(group_index, {})
        previous_result["problem_questions"] = problem_questions
        previous_result["review_remaining"] = len(problem_questions)

    if db.is_enabled():
        try:
            if session.get("saved_quiz_id") and active.get("question_outcomes"):
                await db.record_question_outcomes(
                    user_id=user_id,
                    quiz_id=int(session["saved_quiz_id"]),
                    outcomes=[
                        {"position": pos, "result": result}
                        for pos, result in active["question_outcomes"].items()
                    ],
                )
            special_mode = active.get("special_mode")
            attempt_mode = special_mode if special_mode in ("weak", "bookmarks") else ("review" if active.get("review_mode") else "private")
            await db.save_attempt(
                user_id=user_id,
                username=session.get("owner_username"),
                full_name=session.get("owner_full_name"),
                quiz_id=session.get("saved_quiz_id"),
                mode=attempt_mode,
                chat_id=chat_id,
                section_index=group_index,
                total=total,
                correct=correct,
                wrong=wrong_answered,
                unanswered=unanswered,
                percent=percent,
                best_streak=best_streak,
            )
        except Exception:
            logging.exception("Could not save private attempt")

    buttons = [
        [
            InlineKeyboardButton(
                "🔄 Qayta ishlash",
                callback_data=f"startgroup:{group_index}:{timer_seconds}",
            )
        ],
        [
            InlineKeyboardButton(
                "⏱ Boshqa vaqt bilan",
                callback_data=f"group:{group_index}",
            )
        ],
        [InlineKeyboardButton("📚 Guruhlar", callback_data="groups")],
    ]

    if active.get("special_mode") in ("weak", "bookmarks") and session.get("saved_quiz_id"):
        buttons.insert(0, [InlineKeyboardButton("📘 To‘liq testga qaytish", callback_data=f"pquiz:{int(session['saved_quiz_id'])}")])

    if problem_questions:
        buttons.insert(
            0,
            [
                InlineKeyboardButton(
                    f"❌ {len(problem_questions)} ta xatoni mashq qilish",
                    callback_data=f"retrywrong:{group_index}:{timer_seconds}",
                )
            ],
        )

    if group_index + 1 < len(session["groups"]) and not active.get("review_mode"):
        buttons.insert(
            0,
            [
                InlineKeyboardButton(
                    "➡️ Keyingi guruh",
                    callback_data=f"group:{group_index + 1}",
                )
            ],
        )

    was_review = active.get("review_mode", False)
    session["active"] = None

    if percent == 100:
        try:
            await context.bot.send_message(chat_id=chat_id, text="🏆")
        except Exception:
            pass
    elif percent >= 90:
        try:
            await context.bot.send_message(chat_id=chat_id, text="🎉")
        except Exception:
            pass

    if stopped_reason == "manual":
        title = "🛑 Quiz to‘xtatildi"
    else:
        title = "🏁 Xatolar mashqi tugadi!" if was_review else f"🏁 {group_index + 1}-guruh tugadi!"

    await context.bot.send_message(
        chat_id=chat_id,
        text=(
            f"{title}\n\n"
            f"✅ To‘g‘ri: {correct}\n"
            f"❌ Noto‘g‘ri: {wrong_answered}\n"
            f"⏱ Javobsiz: {unanswered}\n"
            f"🎯 Natija: {percent}%\n"
            f"🔥 Eng yaxshi seriya: {best_streak}\n"
            f"⏲ Vaqt: {format_duration(timer_seconds)}/savol"
        ),
        reply_markup=InlineKeyboardMarkup(buttons),
    )


def group_controller_ok(update: Update, session: dict) -> bool:
    user = update.effective_user
    return bool(user and user.id == session.get("controller_id"))


async def group_choose_size_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = update.effective_chat.id
    session = GROUP_DATA.get(chat_id)

    if not session:
        await query.answer()
        await query.message.reply_text("❌ Guruh sessiyasi topilmadi. Testni qayta yuboring.")
        return
    if not group_controller_ok(update, session):
        await query.answer("Bu quizni faqat boshqaruvchi sozlaydi.", show_alert=True)
        return
    await query.answer()

    size = int(query.data.split(":")[1])
    session.pop("custom_questions", None)
    session.pop("custom_range_label", None)
    session["group_size"] = size
    session.setdefault("shuffle_questions", False)
    session.setdefault("shuffle_options", False)
    touch_group_host(session)
    await send_order_settings(query.message, session, group_mode=True)


async def group_toggle_questions_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    session = GROUP_DATA.get(update.effective_chat.id)
    if not session or not group_controller_ok(update, session):
        await query.answer("Faqat boshqaruvchi sozlaydi.", show_alert=True)
        return
    session["shuffle_questions"] = not bool(session.get("shuffle_questions"))
    touch_group_host(session)
    await query.answer()
    await query.edit_message_reply_markup(
        reply_markup=order_settings_keyboard(session, group_mode=True)
    )


async def group_toggle_options_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    session = GROUP_DATA.get(update.effective_chat.id)
    if not session or not group_controller_ok(update, session):
        await query.answer("Faqat boshqaruvchi sozlaydi.", show_alert=True)
        return
    session["shuffle_options"] = not bool(session.get("shuffle_options"))
    touch_group_host(session)
    await query.answer()
    await query.edit_message_reply_markup(
        reply_markup=order_settings_keyboard(session, group_mode=True)
    )


async def group_toggle_mode_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    session = GROUP_DATA.get(update.effective_chat.id)
    if not session or not group_controller_ok(update, session):
        await query.answer("Faqat boshqaruvchi o‘zgartira oladi.", show_alert=True)
        return
    await query.answer()
    session["quiz_mode"] = "exam" if session.get("quiz_mode", "practice") == "practice" else "practice"
    touch_group_host(session)
    await query.edit_message_reply_markup(reply_markup=order_settings_keyboard(session, True))


async def group_order_done_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = update.effective_chat.id
    session = GROUP_DATA.get(chat_id)
    if not session or not group_controller_ok(update, session):
        await query.answer("Faqat boshqaruvchi sozlaydi.", show_alert=True)
        return
    if not session.get("group_size"):
        await query.answer("Avval bo‘lim hajmini tanlang.", show_alert=True)
        return
    ordered = apply_question_order(session)
    session["ordered_questions"] = ordered
    session["groups"] = build_groups(ordered, session["group_size"])
    touch_group_host(session)
    await query.answer()
    await show_group_quiz_groups(query.message, chat_id)


async def show_group_quiz_groups(message, chat_id: int):
    session = GROUP_DATA.get(chat_id)
    if not session:
        await message.reply_text("❌ Guruh sessiyasi topilmadi.")
        return

    rows = []
    for idx, group in enumerate(session.get("groups", [])):
        start_no = idx * session["group_size"] + 1
        end_no = start_no + len(group) - 1
        if session.get("custom_range_label") and len(session.get("groups", [])) == 1:
            label = f"🎯 {session['custom_range_label']} · {len(group)} ta"
        elif session.get("shuffle_questions"):
            label = f"📘 {idx + 1}-bo‘lim · {len(group)} ta 🔀"
        else:
            label = f"📘 {idx + 1}-bo‘lim · {start_no}-{end_no}"
        rows.append(
            [
                InlineKeyboardButton(
                    label,
                    callback_data=f"ggroup:{idx}",
                )
            ]
        )

    rows.append([InlineKeyboardButton("🏠 Guruh menyusi", callback_data="g_home")])

    await message.reply_text(
        f"✅ {len(session['groups'])} ta bo‘lim tayyor.\n\n"
        "Quiz bo‘limini tanlang:",
        reply_markup=InlineKeyboardMarkup(rows),
    )


async def group_quiz_group_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = update.effective_chat.id
    session = GROUP_DATA.get(chat_id)

    if not session:
        await query.answer()
        await query.message.reply_text("❌ Guruh sessiyasi topilmadi.")
        return
    if not group_controller_ok(update, session):
        await query.answer("Bu testni faqat uni yuklagan foydalanuvchi boshqaradi.", show_alert=True)
        return
    await query.answer()

    group_index = int(query.data.split(":")[1])
    if group_index < 0 or group_index >= len(session.get("groups", [])):
        return

    group = session["groups"][group_index]
    start_no = group_index * session["group_size"] + 1
    end_no = start_no + len(group) - 1

    keyboard = timer_keyboard(group_index, group_mode=True, default_timer=session.get("default_timer"))

    touch_group_host(session)
    range_text = (
        f"🎯 Savollar: {session['custom_range_label']}"
        if session.get("custom_range_label") and len(session.get("groups", [])) == 1
        else ("🔀 Aralashtirilgan savollar" if session.get("shuffle_questions") else f"Savollar: {start_no}-{end_no}")
    )
    await query.message.reply_text(
        f"📘 {group_index + 1}-bo‘lim\n"
        f"{range_text}\n"
        f"Jami: {len(group)}\n\n"
        "⏱ Har bir savol uchun vaqtni tanlang:",
        reply_markup=keyboard,
    )


async def group_timer_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = update.effective_chat.id
    session = GROUP_DATA.get(chat_id)

    if not session:
        await query.answer()
        await query.message.reply_text("❌ Guruh sessiyasi topilmadi.")
        return
    if not group_controller_ok(update, session):
        await query.answer("Bu testni faqat uni yuklagan foydalanuvchi boshqaradi.", show_alert=True)
        return
    await query.answer()

    _, group_text, seconds_text = query.data.split(":")
    group_index = int(group_text)
    seconds = int(seconds_text)

    if seconds not in TIMER_CHOICES:
        return
    if group_index < 0 or group_index >= len(session.get("groups", [])):
        return

    touch_group_host(session)
    session["quiz_mode"] = None
    session["pending_start"] = {
        "group_index": group_index,
        "timer_seconds": seconds,
        "mode": None,
    }

    await query.message.reply_text(
        "🎮 GURUH TESTI REJIMINI TANLANG\n\n"
        "📖 Mashq rejimi\n"
        "• To‘g‘ri/noto‘g‘ri javob savolning o‘zida ko‘rinadi.\n\n"
        "📝 Imtihon rejimi\n"
        "• To‘g‘ri javob savol vaqtida ko‘rsatilmaydi.\n"
        "• Natijalar test oxirida hisoblanadi.\n\n"
        "Rejimni faqat boshqaruvchi tanlaydi:",
        reply_markup=InlineKeyboardMarkup(
            [[
                InlineKeyboardButton(
                    "📖 Mashq rejimi",
                    callback_data=f"gmode:{group_index}:{seconds}:practice",
                ),
                InlineKeyboardButton(
                    "📝 Imtihon rejimi",
                    callback_data=f"gmode:{group_index}:{seconds}:exam",
                ),
            ],[
                InlineKeyboardButton(
                    "⏱ Vaqtni o‘zgartirish",
                    callback_data=f"ggroup:{group_index}",
                )
            ]]
        ),
    )


async def group_mode_choice_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    chat_id = update.effective_chat.id
    session = GROUP_DATA.get(chat_id)

    if not session:
        await query.answer()
        await query.message.reply_text("❌ Guruh sessiyasi topilmadi.")
        return
    if not group_controller_ok(update, session):
        await query.answer("Rejimni faqat boshqaruvchi tanlaydi.", show_alert=True)
        return
    await query.answer()

    _, group_text, seconds_text, mode = query.data.split(":")
    group_index = int(group_text)
    seconds = int(seconds_text)

    if mode not in ("practice", "exam"):
        return
    if seconds not in TIMER_CHOICES:
        return
    if group_index < 0 or group_index >= len(session.get("groups", [])):
        return

    touch_group_host(session)
    session["quiz_mode"] = mode
    session["pending_start"] = {
        "group_index": group_index,
        "timer_seconds": seconds,
        "mode": mode,
    }

    group = session["groups"][group_index]
    max_seconds = len(group) * seconds
    max_time_text = (
        f"{max_seconds / 60:.1f} daqiqagacha"
        if max_seconds >= 60
        else f"{max_seconds} soniyagacha"
    )
    mode_text = "📖 Mashq rejimi" if mode == "practice" else "📝 Imtihon rejimi"

    await query.message.reply_text(
        f"✅ Guruh testi tayyor.\n\n"
        f"📘 Bo‘lim: {group_index + 1}\n"
        f"❓ Savollar: {len(group)}\n"
        f"⏱ Har bir savol: {format_duration(seconds)}\n"
        f"🎮 Rejim: {mode_text}\n"
        f"⌛ Maksimal vaqt: {max_time_text}\n\n"
        "Hamma tayyor bo‘lganda ▶️ START ni bosing.",
        reply_markup=InlineKeyboardMarkup(
            [
                [InlineKeyboardButton(
                    "▶️ START",
                    callback_data=f"gstart:{group_index}:{seconds}",
                )],
                [InlineKeyboardButton(
                    "🎮 Rejimni o‘zgartirish",
                    callback_data=f"gtimer:{group_index}:{seconds}",
                )],
                [InlineKeyboardButton(
                    "⏱ Vaqtni o‘zgartirish",
                    callback_data=f"ggroup:{group_index}",
                )],
                [InlineKeyboardButton("📚 Bo‘limlar", callback_data="ggroups")],
            ]
        ),
    )


async def group_groups_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    session = GROUP_DATA.get(update.effective_chat.id)
    if session and not group_controller_ok(update, session):
        await query.answer("Bu testni faqat uni yuklagan foydalanuvchi boshqaradi.", show_alert=True)
        return
    await query.answer()
    await show_group_quiz_groups(query.message, update.effective_chat.id)


async def group_start_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query

    chat_id = update.effective_chat.id
    session = GROUP_DATA.get(chat_id)

    if not session:
        await query.answer()
        await query.message.reply_text("❌ Guruh sessiyasi topilmadi.")
        return
    if session.get("active"):
        await query.answer("Quiz allaqachon ishlayapti.", show_alert=True)
        return
    if not group_controller_ok(update, session):
        await query.answer("START ni faqat testni yuklagan foydalanuvchi bosishi mumkin.", show_alert=True)
        return
    await query.answer()
    touch_group_host(session)

    _, group_text, seconds_text = query.data.split(":")
    group_index = int(group_text)
    timer_seconds = int(seconds_text)

    if timer_seconds not in TIMER_CHOICES:
        return

    pending = session.get("pending_start") or {}
    if (
        pending.get("group_index") != group_index
        or pending.get("timer_seconds") != timer_seconds
        or pending.get("mode") not in ("practice", "exam")
    ):
        await query.message.reply_text("🎮 Avval Mashq yoki Imtihon rejimini tanlang.")
        return
    session["quiz_mode"] = pending["mode"]
    if group_index < 0 or group_index >= len(session.get("groups", [])):
        return

    session["run_counter"] = session.get("run_counter", 0) + 1
    run_id = session["run_counter"]

    session["active"] = {
        "run_id": run_id,
        "group_index": group_index,
        "questions": session["groups"][group_index],
        "current": 0,
        "timer_seconds": timer_seconds,
        "participants": {},
        "answered_users": set(),
        "empty_streak": 0,
        "paused": False,
        "current_answers": {},
        "current_poll_id": None,
        "current_poll_message_id": None,
        "quiz_mode": session.get("quiz_mode", "practice"),
    }

    try:
        await context.bot.send_dice(chat_id=chat_id, emoji="🎯")
    except Exception:
        pass

    await query.message.reply_text(
        f"🚀 Guruh quizi boshlandi!\n"
        f"📘 {group_index + 1}-bo‘lim\n"
        f"❓ {len(session['active']['questions'])} ta savol\n"
        f"⏱ {format_duration(timer_seconds)}/savol\n"
        f"🔀 Savollar: {'yoqilgan' if session.get('shuffle_questions') else 'o‘chirilgan'} · "
        f"Variantlar: {'yoqilgan' if session.get('shuffle_options') else 'o‘chirilgan'}\n"
        f"🎮 Rejim: {'Imtihon' if session.get('quiz_mode') == 'exam' else 'Mashq'}\n\n"
        "Javob bergan bo‘lsangiz ham, keyingi savol taymer tugagach chiqadi."
    )

    await send_next_group_question(chat_id, context)


async def send_next_group_question(chat_id: int, context: ContextTypes.DEFAULT_TYPE):
    session = GROUP_DATA.get(chat_id)
    if not session or not session.get("active"):
        return

    active = session["active"]

    if active.get("paused"):
        return

    idx = active["current"]
    questions = active["questions"]

    if idx >= len(questions):
        await finish_group_quiz(chat_id, context)
        return

    item = questions[idx]
    displayed_options, displayed_correct_index = prepare_poll_options(
        item,
        bool(session.get("shuffle_options")),
    )
    options = [telegram_safe_option(x) for x in displayed_options]
    timer_seconds = active["timer_seconds"]
    active["answered_users"] = set()
    active["current_answers"] = {}

    try:
        poll_kwargs = dict(
            chat_id=chat_id,
            question=telegram_safe_question(f"[{idx + 1}/{len(questions)}] {item['question']}"),
            options=options,
            is_anonymous=False,
            allows_multiple_answers=False,
            open_period=timer_seconds,
        )
        if active.get("quiz_mode", "practice") == "exam":
            poll_kwargs["type"] = "regular"
        else:
            poll_kwargs["type"] = "quiz"
            poll_kwargs["correct_option_id"] = displayed_correct_index
        msg = await context.bot.send_poll(**poll_kwargs)
    except Exception:
        logging.exception("Could not send group poll")
        await context.bot.send_message(
            chat_id=chat_id,
            text="❌ Bu savol Telegram formatiga sig‘madi. Keyingi savolga o‘tyapman.",
        )
        active["current"] += 1
        await send_next_group_question(chat_id, context)
        return

    active["current_poll_id"] = msg.poll.id
    active["current_poll_message_id"] = msg.message_id

    POLL_MAP[msg.poll.id] = {
        "mode": "group",
        "chat_id": chat_id,
        "message_id": msg.message_id,
        "group_index": active["group_index"],
        "question_index": idx,
        "run_id": active["run_id"],
        "poll_correct_index": displayed_correct_index,
        "handled": False,
    }

    asyncio.create_task(
        group_question_timeout(
            poll_id=msg.poll.id,
            chat_id=chat_id,
            group_index=active["group_index"],
            question_index=idx,
            run_id=active["run_id"],
            timer_seconds=timer_seconds,
            context=context,
        )
    )


async def group_poll_answer_handler(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    meta: dict,
):
    answer = update.poll_answer
    chat_id = meta["chat_id"]
    session = GROUP_DATA.get(chat_id)

    if not session or not session.get("active"):
        return

    active = session["active"]
    if (
        active.get("run_id") != meta.get("run_id")
        or active.get("group_index") != meta.get("group_index")
        or active.get("current") != meta.get("question_index")
        or meta.get("handled")
        or active.get("paused")
    ):
        return

    user = answer.user
    if not user or user.is_bot:
        return

    if user.id in active["answered_users"]:
        return

    selected = answer.option_ids[0] if answer.option_ids else None
    if selected is None:
        return

    active["answered_users"].add(user.id)
    item = active["questions"][meta["question_index"]]
    correct_index = int(meta.get("poll_correct_index", item["correct_index"]))

    # Do not commit group scores until the question timer finishes.
    # This lets /pause safely restart the current question without double-counting.
    active["current_answers"][user.id] = {
        "name": user.full_name or (f"@{user.username}" if user.username else str(user.id)),
        "username": user.username,
        "full_name": user.full_name,
        "is_correct": selected == correct_index,
    }


def commit_group_current_answers(active: dict) -> int:
    answers = active.get("current_answers", {})
    for user_id, answer_data in answers.items():
        participant = active["participants"].setdefault(
            user_id,
            {
                "name": answer_data.get("name") or str(user_id),
                "username": answer_data.get("username"),
                "full_name": answer_data.get("full_name"),
                "correct": 0,
                "wrong": 0,
                "current_streak": 0,
                "best_streak": 0,
            },
        )

        if answer_data.get("is_correct"):
            participant["correct"] += 1
            participant["current_streak"] += 1
            participant["best_streak"] = max(
                participant["best_streak"],
                participant["current_streak"],
            )
        else:
            participant["wrong"] += 1
            participant["current_streak"] = 0

    count = len(answers)
    active["current_answers"] = {}
    active["answered_users"] = set()
    return count


async def group_question_timeout(
    poll_id: str,
    chat_id: int,
    group_index: int,
    question_index: int,
    run_id: int,
    timer_seconds: int,
    context: ContextTypes.DEFAULT_TYPE,
):
    await asyncio.sleep(timer_seconds + 0.8)

    meta = POLL_MAP.get(poll_id)
    if not meta or meta.get("handled"):
        return

    session = GROUP_DATA.get(chat_id)
    if not session or not session.get("active"):
        POLL_MAP.pop(poll_id, None)
        return

    active = session["active"]
    if (
        active.get("run_id") != run_id
        or active.get("group_index") != group_index
        or active.get("current") != question_index
    ):
        POLL_MAP.pop(poll_id, None)
        return

    meta["handled"] = True

    answered_count = commit_group_current_answers(active)
    if answered_count == 0:
        active["empty_streak"] = active.get("empty_streak", 0) + 1
    else:
        active["empty_streak"] = 0

    active["current"] += 1
    active["current_poll_id"] = None
    active["current_poll_message_id"] = None
    POLL_MAP.pop(poll_id, None)

    if active["empty_streak"] >= GROUP_EMPTY_STOP_THRESHOLD:
        await context.bot.send_message(
            chat_id=chat_id,
            text=(
                f"😴 Oxirgi {GROUP_EMPTY_STOP_THRESHOLD} ta savolga hech kim javob bermadi.\n"
                "Quiz avtomatik to‘xtatildi va hozirgacha bo‘lgan natija chiqariladi."
            ),
        )
        await finish_group_quiz(
            chat_id,
            context,
            completed_count=active["current"],
            stopped_reason="no_answers",
        )
        return

    await asyncio.sleep(QUESTION_TRANSITION_DELAY)
    await send_next_group_question(chat_id, context)


async def finish_group_quiz(chat_id: int, context: ContextTypes.DEFAULT_TYPE, completed_count: Optional[int] = None, stopped_reason: Optional[str] = None):
    session = GROUP_DATA.get(chat_id)
    if not session or not session.get("active"):
        return

    active = session["active"]
    total = completed_count if completed_count is not None else len(active["questions"])
    group_index = active["group_index"]
    timer_seconds = active["timer_seconds"]
    participants = active["participants"]

    ranking = []
    for user_id, p in participants.items():
        unanswered = max(0, total - p["correct"] - p["wrong"])
        percent = round((p["correct"] / total) * 100) if total else 0
        ranking.append(
            {
                "user_id": user_id,
                "name": p["name"],
                "correct": p["correct"],
                "wrong": p["wrong"],
                "unanswered": unanswered,
                "percent": percent,
                "best_streak": p["best_streak"],
            }
        )

    if db.is_enabled():
        for row in ranking:
            p = participants.get(row["user_id"], {})
            try:
                await db.save_attempt(
                    user_id=row["user_id"],
                    username=p.get("username"),
                    full_name=p.get("full_name") or row.get("name"),
                    quiz_id=session.get("saved_quiz_id"),
                    mode="group",
                    chat_id=chat_id,
                    section_index=group_index,
                    total=total,
                    correct=row["correct"],
                    wrong=row["wrong"],
                    unanswered=row["unanswered"],
                    percent=row["percent"],
                    best_streak=row["best_streak"],
                )
            except Exception:
                logging.exception("Could not save group attempt for %s", row["user_id"])

    ranking.sort(
        key=lambda x: (
            -x["correct"],
            x["wrong"],
            x["unanswered"],
            x["name"].lower(),
        )
    )

    finish_title = "🛑 GURUH QUIZI TO‘XTATILDI" if stopped_reason else "🏆 GURUH QUIZI TUGADI"
    lines = [
        finish_title,
        "",
        f"📘 Bo‘lim: {group_index + 1}",
        f"❓ Savollar: {total}",
        f"👥 Qatnashchilar: {len(ranking)}",
    ]
    if stopped_reason == "manual":
        lines.append("🛑 Boshqaruvchi quizni to‘xtatdi.")
    elif stopped_reason == "no_answers":
        lines.append("😴 Faollik bo‘lmagani uchun avtomatik to‘xtadi.")
    lines.extend(["", "🏅 Reyting:"])

    medals = ["🥇", "🥈", "🥉"]
    for pos, row in enumerate(ranking[:20], start=1):
        medal = medals[pos - 1] if pos <= 3 else f"{pos}."
        lines.append(
            f"{medal} {row['name']} — {row['correct']}/{total} "
            f"({row['percent']}%) · 🔥 {row['best_streak']}"
        )

    if not ranking:
        lines.append("Hech kim javob bermadi.")

    leaderboard_text = "\n".join(lines)
    session["last_leaderboard_text"] = leaderboard_text
    session.setdefault("group_results", {})[group_index] = ranking
    session["active"] = None
    touch_group_host(session)

    try:
        await context.bot.send_message(chat_id=chat_id, text="🎉")
    except Exception:
        pass

    buttons = [
        [
            InlineKeyboardButton(
                "🔄 Qayta boshlash",
                callback_data=f"gstart:{group_index}:{timer_seconds}",
            )
        ],
        [InlineKeyboardButton("📚 Bo‘limlar", callback_data="ggroups")],
        [InlineKeyboardButton("🔓 Sessiyani tugatish", callback_data="g_release")],
        [InlineKeyboardButton("🏠 Guruh menyusi", callback_data="g_home")],
    ]

    if group_index + 1 < len(session.get("groups", [])):
        buttons.insert(
            0,
            [
                InlineKeyboardButton(
                    "➡️ Keyingi bo‘lim",
                    callback_data=f"ggroup:{group_index + 1}",
                )
            ],
        )

    await context.bot.send_message(
        chat_id=chat_id,
        text=leaderboard_text,
        reply_markup=InlineKeyboardMarkup(buttons),
    )


async def stop_group_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Stop the active quiz completely in private chat or group."""
    chat_id = update.effective_chat.id

    if not is_group_chat(update.effective_chat):
        user_id = update.effective_user.id
        session = USER_DATA.get(user_id)
        if not session or not session.get("active"):
            await update.message.reply_text("🛑 Hozir faol quiz yo‘q.")
            return

        active = session["active"]
        poll_id = active.get("current_poll_id")
        message_id = active.get("current_poll_message_id")

        if poll_id:
            meta = POLL_MAP.get(poll_id)
            if meta:
                meta["handled"] = True
            POLL_MAP.pop(poll_id, None)

        if message_id:
            try:
                current_item = (
                    active["questions"][active["current"]]
                    if 0 <= active.get("current", -1) < len(active.get("questions", []))
                    else None
                )
                markup = private_bookmark_markup(session, active, current_item) if current_item else None
                await context.bot.stop_poll(chat_id, message_id, reply_markup=markup)
            except Exception:
                try:
                    await context.bot.stop_poll(chat_id, message_id)
                except Exception:
                    pass

        active["current_poll_id"] = None
        active["current_poll_message_id"] = None
        completed = min(active["current"], len(active["questions"]))
        await finish_group(
            chat_id,
            user_id,
            context,
            completed_count=completed,
            stopped_reason="manual",
        )
        return

    session = GROUP_DATA.get(chat_id)
    if not session or not session.get("active"):
        await update.message.reply_text("🛑 Hozir faol guruh quizi yo‘q.")
        return
    if update.effective_user.id != session.get("controller_id"):
        await update.message.reply_text("⛔ Quizni faqat uni boshlagan boshqaruvchi to‘xtata oladi.")
        return

    active = session["active"]
    poll_id = active.get("current_poll_id")
    message_id = active.get("current_poll_message_id")

    if poll_id:
        meta = POLL_MAP.get(poll_id)
        if meta:
            meta["handled"] = True
        POLL_MAP.pop(poll_id, None)

    if message_id:
        try:
            await context.bot.stop_poll(chat_id, message_id)
        except Exception:
            pass

    # Keep answers already submitted on the visible question before stopping.
    commit_group_current_answers(active)
    completed = min(active["current"] + (1 if message_id else 0), len(active["questions"]))
    active["current_poll_id"] = None
    active["current_poll_message_id"] = None

    await finish_group_quiz(
        chat_id,
        context,
        completed_count=completed,
        stopped_reason="manual",
    )


async def pause_quiz_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Temporarily pause the current quiz. /resume continues it."""
    chat_id = update.effective_chat.id

    if not is_group_chat(update.effective_chat):
        user_id = update.effective_user.id
        session = USER_DATA.get(user_id)
        if not session or not session.get("active"):
            await update.message.reply_text("⏸ Hozir faol quiz yo‘q.")
            return

        active = session["active"]
        if active.get("paused"):
            await update.message.reply_text("⏸ Quiz allaqachon pauzada. /resume bilan davom eting.")
            return

        poll_id = active.get("current_poll_id")
        message_id = active.get("current_poll_message_id")
        if poll_id:
            meta = POLL_MAP.get(poll_id)
            if meta:
                meta["handled"] = True
            POLL_MAP.pop(poll_id, None)
        if message_id:
            try:
                current_item = (
                    active["questions"][active["current"]]
                    if 0 <= active.get("current", -1) < len(active.get("questions", []))
                    else None
                )
                markup = private_bookmark_markup(session, active, current_item) if current_item else None
                await context.bot.stop_poll(chat_id, message_id, reply_markup=markup)
            except Exception:
                try:
                    await context.bot.stop_poll(chat_id, message_id)
                except Exception:
                    pass

        # The interrupted question is not counted; it will be sent again on resume.
        active["current_poll_id"] = None
        active["current_poll_message_id"] = None
        active["paused"] = True
        await update.message.reply_text(
            "⏸ Quiz pauzaga qo‘yildi. Joriy savol hisoblanmadi.\n"
            "Davom ettirish uchun /resume yuboring."
        )
        return

    session = GROUP_DATA.get(chat_id)
    if not session or not session.get("active"):
        await update.message.reply_text("⏸ Hozir faol guruh quizi yo‘q.")
        return
    if update.effective_user.id != session.get("controller_id"):
        await update.message.reply_text("⛔ Quizni faqat boshqaruvchi pauza qila oladi.")
        return

    active = session["active"]
    if active.get("paused"):
        await update.message.reply_text("⏸ Guruh quizi allaqachon pauzada.")
        return

    poll_id = active.get("current_poll_id")
    message_id = active.get("current_poll_message_id")
    if poll_id:
        meta = POLL_MAP.get(poll_id)
        if meta:
            meta["handled"] = True
        POLL_MAP.pop(poll_id, None)
    if message_id:
        try:
            await context.bot.stop_poll(chat_id, message_id)
        except Exception:
            pass

    # Answers on an interrupted group question are discarded so nobody is double-counted.
    active["current_answers"] = {}
    active["answered_users"] = set()
    active["current_poll_id"] = None
    active["current_poll_message_id"] = None
    active["paused"] = True
    await update.message.reply_text(
        "⏸ Guruh quizi pauzaga qo‘yildi. Joriy savol hisoblanmadi.\n"
        "Davom ettirish uchun /resume yuboring."
    )


async def resume_quiz_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Resume a manually/automatically paused quiz."""
    chat_id = update.effective_chat.id

    if not is_group_chat(update.effective_chat):
        user_id = update.effective_user.id
        session = USER_DATA.get(user_id)
        if not session or not session.get("active"):
            await update.message.reply_text("▶️ Davom ettiriladigan quiz topilmadi.")
            return

        active = session["active"]
        if not active.get("paused"):
            await update.message.reply_text("▶️ Quiz pauzada emas.")
            return

        active["paused"] = False
        active["empty_streak"] = 0
        await update.message.reply_text("▶️ Quiz davom etmoqda.")
        await send_next_question(chat_id, user_id, context)
        return

    session = GROUP_DATA.get(chat_id)
    if not session or not session.get("active"):
        await update.message.reply_text("▶️ Davom ettiriladigan guruh quizi topilmadi.")
        return
    if update.effective_user.id != session.get("controller_id"):
        await update.message.reply_text("⛔ Quizni faqat boshqaruvchi davom ettira oladi.")
        return

    active = session["active"]
    if not active.get("paused"):
        await update.message.reply_text("▶️ Guruh quizi pauzada emas.")
        return

    active["paused"] = False
    active["empty_streak"] = 0
    touch_group_host(session)
    await update.message.reply_text("▶️ Guruh quizi davom etmoqda.")
    await send_next_group_question(chat_id, context)


async def group_stop_no_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer("Davom etadi")
    await query.message.reply_text("▶️ Quiz davom etmoqda.")


async def group_stop_yes_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    session = GROUP_DATA.get(update.effective_chat.id)
    if not session or not session.get("active"):
        await query.answer("Faol quiz yo‘q", show_alert=True)
        return
    if update.effective_user.id != session.get("controller_id"):
        await query.answer("Faqat quiz boshqaruvchisi to‘xtata oladi.", show_alert=True)
        return
    await query.answer()

    active = session["active"]
    poll_id = active.get("current_poll_id")
    message_id = active.get("current_poll_message_id")

    if poll_id:
        meta = POLL_MAP.get(poll_id)
        if meta:
            meta["handled"] = True
        POLL_MAP.pop(poll_id, None)

    if message_id:
        try:
            await context.bot.stop_poll(update.effective_chat.id, message_id)
        except Exception:
            pass

    # Keep submitted answers on the visible question before stopping.
    commit_group_current_answers(active)
    # The currently visible question counts as seen.
    completed = min(active["current"] + (1 if message_id else 0), len(active["questions"]))
    active["current_poll_id"] = None
    active["current_poll_message_id"] = None

    await finish_group_quiz(
        update.effective_chat.id,
        context,
        completed_count=completed,
        stopped_reason="manual",
    )


async def skip_group_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_group_chat(update.effective_chat):
        await update.message.reply_text("/skip faqat guruh quizida ishlaydi.")
        return

    session = GROUP_DATA.get(update.effective_chat.id)
    if not session or not session.get("active"):
        await update.message.reply_text("⏭ Hozir faol guruh quizi yo‘q.")
        return
    if update.effective_user.id != session.get("controller_id"):
        await update.message.reply_text("⛔ Savolni faqat quiz boshqaruvchisi o‘tkaza oladi.")
        return

    active = session["active"]
    poll_id = active.get("current_poll_id")
    message_id = active.get("current_poll_message_id")
    if not poll_id or not message_id:
        await update.message.reply_text("⏭ Hozir o‘tkazib yuboriladigan savol yo‘q.")
        return

    meta = POLL_MAP.get(poll_id)
    if meta:
        meta["handled"] = True
    POLL_MAP.pop(poll_id, None)

    try:
        await context.bot.stop_poll(update.effective_chat.id, message_id)
    except Exception:
        pass

    active["current_answers"] = {}
    active["answered_users"] = set()
    active["current"] += 1
    active["empty_streak"] = 0
    active["current_poll_id"] = None
    active["current_poll_message_id"] = None

    await update.message.reply_text("⏭ Savol o‘tkazib yuborildi.")
    await send_next_group_question(update.effective_chat.id, context)


async def db_status_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    status = "✅ Database ulangan" if db.is_enabled() else "❌ Database ulanmagan"
    await update.message.reply_text(status)


async def app_post_init(application: Application):
    try:
        await application.bot.set_my_commands(
            PRIVATE_COMMAND_MENU,
            scope=BotCommandScopeDefault(),
        )
        await application.bot.set_my_commands(
            PRIVATE_COMMAND_MENU,
            scope=BotCommandScopeAllPrivateChats(),
        )
        await application.bot.set_my_commands(
            GROUP_COMMAND_MENU,
            scope=BotCommandScopeAllGroupChats(),
        )
        logging.info("Telegram command menus configured.")
    except Exception:
        logging.exception("Could not configure Telegram command menus")

    try:
        await db.init_pool()
    except Exception:
        logging.exception("Database initialization failed")


async def app_post_shutdown(application: Application):
    try:
        await db.close_pool()
    except Exception:
        logging.exception("Database shutdown failed")


async def retry_wrong_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = update.effective_user.id
    chat_id = update.effective_chat.id
    session = USER_DATA.get(user_id)

    if not session:
        await query.message.reply_text("❌ Sessiya topilmadi.")
        return

    _, group_text, timer_text = query.data.split(":")
    group_index = int(group_text)
    timer_seconds = int(timer_text)

    result = session.get("results", {}).get(group_index)
    problem_questions = (result or {}).get("problem_questions", [])

    if not problem_questions:
        await query.message.reply_text("✅ Mashq qilish uchun xato savollar qolmagan.")
        return

    # Remove the used "practice mistakes" button from the old result card.
    # The new review result will show a fresh button only for mistakes that remain.
    try:
        if query.message.reply_markup:
            cleaned_rows = []
            for row in query.message.reply_markup.inline_keyboard:
                kept = [
                    button for button in row
                    if not (button.callback_data or "").startswith("retrywrong:")
                ]
                if kept:
                    cleaned_rows.append(kept)
            await query.edit_message_reply_markup(
                reply_markup=InlineKeyboardMarkup(cleaned_rows) if cleaned_rows else None
            )
    except Exception:
        logging.exception("Could not remove stale retry button")

    session["run_counter"] = session.get("run_counter", 0) + 1
    run_id = session["run_counter"]

    session["active"] = {
        "run_id": run_id,
        "group_index": group_index,
        "questions": problem_questions,
        "current": 0,
        "correct": 0,
        "wrong": [],
        "unanswered": [],
        "answered_polls": set(),
        "timer_seconds": timer_seconds,
        "current_streak": 0,
        "best_streak": 0,
        "review_mode": True,
        "empty_streak": 0,
        "paused": False,
        "current_poll_id": None,
        "current_poll_message_id": None,
        "quiz_mode": "practice",
        "question_outcomes": {},
        "special_mode": "review",
        "pro_features": await user_has_pro(user_id),
        "bookmarked_positions": set(),
    }
    if session["active"]["pro_features"] and session.get("saved_quiz_id") and db.is_enabled():
        try:
            session["active"]["bookmarked_positions"] = set(
                await db.list_bookmark_positions(user_id, int(session["saved_quiz_id"]))
            )
        except Exception:
            logging.exception("Bookmarklar yuklanmadi")

    try:
        await context.bot.send_dice(chat_id=chat_id, emoji="🎯")
    except Exception:
        pass

    await query.message.reply_text(
        f"🧠 Xatolar mashqi boshlandi!\n"
        f"❓ {len(problem_questions)} ta savol\n"
        f"⏱ {format_duration(timer_seconds)}/savol"
    )
    await send_next_question(chat_id, user_id, context)


async def global_error_handler(update: object, context: ContextTypes.DEFAULT_TYPE):
    err = context.error
    if err:
        logging.error("Botda kutilmagan xato", exc_info=(type(err), err, err.__traceback__))
    try:
        chat = getattr(update, "effective_chat", None)
        if chat:
            await context.bot.send_message(
                chat_id=chat.id,
                text="⚠️ Texnik xato yuz berdi. Saqlangan testlaringiz bazada qoladi. /start orqali bosh menyuga qayting.",
            )
    except Exception:
        logging.exception("Xato xabarini yuborib bo‘lmadi")


def build_telegram_application() -> Application:
    """Create the Telegram application without long polling."""
    if not TELEGRAM_TOKEN:
        raise RuntimeError("TELEGRAM_TOKEN environment variable is missing.")

    app = Application.builder().token(TELEGRAM_TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("new", new_command))
    app.add_handler(CommandHandler("quizzes", quizzes_command))
    app.add_handler(CommandHandler("continue", continue_command))
    app.add_handler(CommandHandler("progress", progress_command))
    app.add_handler(CommandHandler("group", group_mode_command))
    app.add_handler(CommandHandler("settings", settings_command))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("cancel", cancel_command))
    app.add_handler(CommandHandler("pause", pause_quiz_command))
    app.add_handler(CommandHandler("resume", resume_quiz_command))
    app.add_handler(CommandHandler("stop", stop_group_command))
    app.add_handler(CommandHandler("skip", skip_group_command))
    app.add_handler(CommandHandler("release", release_group_command))
    app.add_handler(CommandHandler("dbstatus", db_status_command))
    app.add_handler(CommandHandler("parser", parser_command))
    app.add_handler(CommandHandler("plan", plan_command))
    app.add_handler(CommandHandler("myid", myid_command))
    app.add_handler(CommandHandler("grantpro", grantpro_command))
    app.add_handler(CommandHandler("revokepro", revokepro_command))
    app.add_handler(
        MessageHandler(
            filters.Document.PDF
            | filters.Document.FileExtension("docx"),
            handle_document,
        )
    )
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_plain_text))

    app.add_handler(CallbackQueryHandler(home_callback, pattern=r"^menu_home$"))
    app.add_handler(CallbackQueryHandler(group_home_callback, pattern=r"^g_home$"))
    app.add_handler(CallbackQueryHandler(group_release_callback, pattern=r"^g_release$"))
    app.add_handler(CallbackQueryHandler(group_parser_callback, pattern=r"^g_parser$"))
    app.add_handler(CallbackQueryHandler(group_saved_callback, pattern=r"^g_saved$"))
    app.add_handler(CallbackQueryHandler(private_quiz_detail_callback, pattern=r"^pquiz:\d+$"))
    app.add_handler(CallbackQueryHandler(private_load_saved_callback, pattern=r"^pload:\d+$"))
    app.add_handler(CallbackQueryHandler(private_duplicate_quiz_callback, pattern=r"^pduplicate:\d+$"))
    app.add_handler(CallbackQueryHandler(private_reset_stats_callback, pattern=r"^presetstats:\d+$"))
    app.add_handler(CallbackQueryHandler(private_reset_stats_yes_callback, pattern=r"^presetstatsyes:\d+$"))
    app.add_handler(CallbackQueryHandler(private_rename_quiz_callback, pattern=r"^prename:\d+$"))
    app.add_handler(CallbackQueryHandler(private_delete_quiz_callback, pattern=r"^pdelete:\d+$"))
    app.add_handler(CallbackQueryHandler(private_delete_quiz_yes_callback, pattern=r"^pdeleteyes:\d+$"))
    app.add_handler(CallbackQueryHandler(private_quiz_results_callback, pattern=r"^presults:\d+$"))
    app.add_handler(CallbackQueryHandler(weak_questions_callback, pattern=r"^pweak:\d+$"))
    app.add_handler(CallbackQueryHandler(weak_start_callback, pattern=r"^pweakstart:\d+$"))
    app.add_handler(CallbackQueryHandler(bookmarks_callback, pattern=r"^pbookmarks:\d+$"))
    app.add_handler(CallbackQueryHandler(bookmarks_start_callback, pattern=r"^pbookstart:\d+$"))
    app.add_handler(CallbackQueryHandler(bookmark_toggle_callback, pattern=r"^pbookmarktoggle:\d+:\d+:\d+$"))
    app.add_handler(CallbackQueryHandler(group_load_saved_callback, pattern=r"^gload:\d+$"))
    app.add_handler(CallbackQueryHandler(group_stop_yes_callback, pattern=r"^gstop_yes$"))
    app.add_handler(CallbackQueryHandler(group_stop_no_callback, pattern=r"^gstop_no$"))
    app.add_handler(CallbackQueryHandler(group_new_callback, pattern=r"^g_new$"))
    app.add_handler(CallbackQueryHandler(group_current_callback, pattern=r"^g_current$"))
    app.add_handler(CallbackQueryHandler(group_leaderboard_callback, pattern=r"^g_leaderboard$"))
    app.add_handler(CallbackQueryHandler(group_help_callback, pattern=r"^g_help$"))
    app.add_handler(CallbackQueryHandler(new_callback, pattern=r"^menu_new$"))
    app.add_handler(CallbackQueryHandler(quizzes_callback, pattern=r"^menu_quizzes$"))
    app.add_handler(CallbackQueryHandler(continue_callback, pattern=r"^menu_continue$"))
    app.add_handler(CallbackQueryHandler(progress_callback, pattern=r"^menu_progress$"))
    app.add_handler(CallbackQueryHandler(group_mode_callback, pattern=r"^menu_group$"))
    app.add_handler(CallbackQueryHandler(settings_callback, pattern=r"^menu_settings$"))
    app.add_handler(CallbackQueryHandler(prefs_toggle_callback, pattern=r"^prefs_(?:qshuffle|oshuffle)$"))
    app.add_handler(CallbackQueryHandler(prefs_size_callback, pattern=r"^prefs_size$"))
    app.add_handler(CallbackQueryHandler(prefs_set_size_callback, pattern=r"^prefs_set_size:(?:30|40|50|100)$"))
    app.add_handler(CallbackQueryHandler(prefs_timer_callback, pattern=r"^prefs_timer$"))
    app.add_handler(CallbackQueryHandler(prefs_set_timer_callback, pattern=r"^prefs_set_timer:(?:10|15|20|30|40|60|120)$"))
    app.add_handler(CallbackQueryHandler(pro_info_callback, pattern=r"^pro_info$"))
    app.add_handler(CallbackQueryHandler(help_callback, pattern=r"^menu_help$"))
    app.add_handler(CallbackQueryHandler(help_formats_callback, pattern=r"^help_formats$"))
    app.add_handler(CallbackQueryHandler(help_ai_callback, pattern=r"^help_ai$"))
    app.add_handler(CallbackQueryHandler(help_parser_info_callback, pattern=r"^help_parser_info$"))
    app.add_handler(CallbackQueryHandler(help_about_callback, pattern=r"^help_about$"))
    app.add_handler(CallbackQueryHandler(help_privacy_callback, pattern=r"^help_privacy$"))
    app.add_handler(CallbackQueryHandler(help_upload_callback, pattern=r"^help_upload$"))
    app.add_handler(CallbackQueryHandler(choose_size_callback, pattern=r"^size:\d+$"))
    app.add_handler(CallbackQueryHandler(custom_range_callback, pattern=r"^(?:p|g)customrange$"))
    app.add_handler(CallbackQueryHandler(custom_random_callback, pattern=r"^(?:p|g)customrandom$"))
    app.add_handler(CallbackQueryHandler(private_toggle_questions_callback, pattern=r"^ptoq$"))
    app.add_handler(CallbackQueryHandler(private_toggle_options_callback, pattern=r"^ptoa$"))
    app.add_handler(CallbackQueryHandler(private_order_done_callback, pattern=r"^porderdone$"))
    app.add_handler(CallbackQueryHandler(group_choose_size_callback, pattern=r"^gsize:\d+$"))
    app.add_handler(CallbackQueryHandler(group_toggle_questions_callback, pattern=r"^gtoq$"))
    app.add_handler(CallbackQueryHandler(group_toggle_options_callback, pattern=r"^gtoa$"))
    app.add_handler(CallbackQueryHandler(group_order_done_callback, pattern=r"^gorderdone$"))
    app.add_handler(CallbackQueryHandler(group_quiz_group_callback, pattern=r"^ggroup:\d+$"))
    app.add_handler(CallbackQueryHandler(group_groups_callback, pattern=r"^ggroups$"))
    app.add_handler(CallbackQueryHandler(group_timer_callback, pattern=r"^gtimer:\d+:(?:10|15|20|30|40|60|120)$"))
    app.add_handler(CallbackQueryHandler(group_mode_choice_callback, pattern=r"^gmode:\d+:(?:10|15|20|30|40|60|120):(?:practice|exam)$"))
    app.add_handler(CallbackQueryHandler(group_start_callback, pattern=r"^gstart:\d+:(?:10|15|20|30|40|60|120)$"))
    app.add_handler(CallbackQueryHandler(retry_wrong_callback, pattern=r"^retrywrong:\d+:(?:10|15|20|30|40|60|120)$"))
    app.add_handler(CallbackQueryHandler(group_callback, pattern=r"^group:\d+$"))
    app.add_handler(CallbackQueryHandler(timer_callback, pattern=r"^timer:\d+:(?:10|15|20|30|40|60|120)$"))
    app.add_handler(CallbackQueryHandler(private_mode_choice_callback, pattern=r"^pmode:\d+:(?:10|15|20|30|40|60|120):(?:practice|exam)$"))
    app.add_handler(CallbackQueryHandler(groups_callback, pattern=r"^groups$"))
    app.add_handler(CallbackQueryHandler(start_group_callback, pattern=r"^startgroup:\d+:(?:10|15|20|30|40|60|120)$"))
    app.add_handler(PollAnswerHandler(poll_answer_handler))

    app.add_error_handler(global_error_handler)
    return app


@asynccontextmanager
async def webhook_lifespan(app: FastAPI):
    """Run Telegram and the health API in one Render web process."""
    global TELEGRAM_APP

    telegram_app = build_telegram_application()
    TELEGRAM_APP = telegram_app

    try:
        await telegram_app.initialize()

        # In FastAPI integration, call our existing initialization hook directly.
        await app_post_init(telegram_app)

        # Starts python-telegram-bot's internal update processor only.
        # It does NOT start getUpdates/long polling.
        await telegram_app.start()

        # Register/refresh the webhook. Once a webhook is active,
        # Telegram does not allow getUpdates for this bot.
        await telegram_app.bot.set_webhook(
            url=WEBHOOK_URL,
            allowed_updates=Update.ALL_TYPES,
            drop_pending_updates=False,
            secret_token=WEBHOOK_SECRET,
        )

        logging.info("Testchi webhook rejimida ishga tushdi: %s", WEBHOOK_URL)
        yield

    finally:
        try:
            if telegram_app.running:
                await telegram_app.stop()
        except Exception:
            logging.exception("Telegram application stop failed")

        try:
            await app_post_shutdown(telegram_app)
        except Exception:
            logging.exception("Application shutdown hook failed")

        try:
            await telegram_app.shutdown()
        except Exception:
            logging.exception("Telegram application shutdown failed")

        TELEGRAM_APP = None


web_app = FastAPI(
    title="Testchi",
    version="1.0",
    lifespan=webhook_lifespan,
)


@web_app.get("/")
async def root():
    return {
        "status": "ok",
        "service": "test-tuzuvchi",
        "telegram_mode": "webhook",
    }


@web_app.get("/health")
async def health():
    return {
        "status": "healthy",
        "telegram_mode": "webhook",
        "telegram_app_running": bool(TELEGRAM_APP and TELEGRAM_APP.running),
        "database_enabled": db.is_enabled(),
    }


@web_app.post(WEBHOOK_PATH)
async def telegram_webhook(request: Request):
    """Receive a Telegram update and queue it for the PTB application."""
    if TELEGRAM_APP is None or not TELEGRAM_APP.running:
        raise HTTPException(status_code=503, detail="Telegram application is starting")

    supplied_secret = request.headers.get(
        "X-Telegram-Bot-Api-Secret-Token", ""
    )
    if supplied_secret != WEBHOOK_SECRET:
        raise HTTPException(status_code=403, detail="Invalid webhook secret")

    try:
        payload = await request.json()
        update = Update.de_json(payload, TELEGRAM_APP.bot)
    except Exception as exc:
        logging.warning("Invalid Telegram webhook payload: %s", exc)
        raise HTTPException(status_code=400, detail="Invalid Telegram update")

    # Return to Telegram quickly; PTB processes the update from its queue.
    await TELEGRAM_APP.update_queue.put(update)
    return {"ok": True}


def main():
    """
    Polling is intentionally disabled.

    The existing Render command can stay:
    python bot.py & uvicorn api:app --host 0.0.0.0 --port $PORT

    `python bot.py` exits immediately; `api:app` is the only live service.
    """
    logging.info(
        "Webhook versiya: polling ishga tushirilmaydi. "
        "Uvicorn api:app webhook serverni boshqaradi."
    )


if __name__ == "__main__":
    main()
